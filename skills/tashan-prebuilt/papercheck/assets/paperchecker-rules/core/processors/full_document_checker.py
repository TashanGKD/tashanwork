import docx
import re
from typing import List, Dict, Se
import os
import json
import requests
import pymupdf4llm
from concurrent.futures import ThreadPoolExecutor, as_completed
import time
import logging
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import random
import hashlib
import sys
import tempfile

# 尝试导入AI增强引用提取模块 - 使用新的模块路径
try:
    from ..extractor.ai_extractor import extract_references, extract_western_references
    AI_ENHANCED_EXTRACTION_AVAILABLE = True
except ImportError:
    AI_ENHANCED_EXTRACTION_AVAILABLE = False
    print("警告: 未找到AI增强引用提取模块")

# 尝试导入Markdown引用提取模块 - 使用新的模块路径
try:
    from ..extractor.ai_extractor import extract_citations_from_text as extract_citations_from_markdown
    MARKDOWN_EXTRACTION_AVAILABLE = True
except ImportError:
    MARKDOWN_EXTRACTION_AVAILABLE = False
    print("警告: 未找到Markdown引用提取模块")

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# 创建缓存目录
CACHE_DIR = "pdf_cache"
if not os.path.exists(CACHE_DIR):
    os.makedirs(CACHE_DIR)

# 定义重试策略
# 对于HTTP 429 (Too Many Requests) 和 网络连接问题进行重试
retry_strategy = retry(
    stop=stop_after_attempt(3),  # 最多重试3次
    wait=wait_exponential(multiplier=1, min=4, max=10),  # 指数退避，最小4秒，最大10秒
    retry=retry_if_exception_type((requests.exceptions.ConnectionError,
                                   requests.exceptions.Timeout,
                                   requests.exceptions.RequestException))
)

class CitationChecker:
    def __init__(self, doc_path: str, config_path: str = "config.json"):
        self.doc_path = doc_path
        self.doc = None
        self.md_content = None
        self.temp_md_file = None  # 临时 Markdown 文件路径

        # 检查是否为PDF文件
        if doc_path.lower().endswith('.pdf'):
            # 处理PDF文件
            from utils.mineru_pdf_converter import convert_pdf_to_markdown
            try:
                # 将PDF转换为Markdown
                self.temp_md_file = convert_pdf_to_markdown(doc_path, config_path=config_path)

                # 读取转换后的Markdown内容
                with open(self.temp_md_file, 'r', encoding='utf-8') as f:
                    self.md_content = f.read()

                # 创建一个模拟的docx.Document对象用于后续处理
                # 我们将模拟段落结构
                self._create_doc_from_markdown()
            except Exception as e:
                print(f"PDF转换失败: {e}")
                raise
        else:
            # 处理Word文档
            self.doc = docx.Document(doc_path)
        self.citations = []  # 存储文中引用
        self.references = []  # 存储参考文献
        self.missing_citations = []  # 存储未在参考文献中出现的引用

        # 读取配置文件
        try:
            from config.config_manager import ConfigManager
            # 确保配置路径正确，处理相对路径和绝对路径

            original_config_path = config_path  # 保存原始路径用于错误消息
            config_manager_initialized = False

            if not os.path.isabs(config_path):
                # 如果是相对路径，尝试在当前目录和config目录中查找
                if not os.path.exists(config_path):
                    # 尝试在项目根目录的config子目录中查找
                    project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
                    config_in_config_dir = os.path.join(project_root, "config", config_path)
                    if os.path.exists(config_in_config_dir):
                        config_path = config_in_config_dir
                    else:
                        # 如果配置文件不存在，使用默认配置
                        print(f"配置文件不存在: {original_config_path}, 使用默认配置")
                        self.config = self._get_default_config(config_path)
                        self.config_manager = None
                        config_manager_initialized = False
                else:
                    config_manager_initialized = True
                    self.config_manager = ConfigManager(config_path)
                    self.config = self.config_manager.get_config()

                # 如果在config目录中找到了文件
                if os.path.exists(config_path) and not config_manager_initialized:
                    config_manager_initialized = True
                    self.config_manager = ConfigManager(config_path)
                    self.config = self.config_manager.get_config()
            else:
                # 绝对路径
                config_manager_initialized = True
                self.config_manager = ConfigManager(config_path)
                self.config = self.config_manager.get_config()

            # 如果config_manager没有初始化，使用默认配置
            if not config_manager_initialized:
                self.config = self._get_default_config(original_config_path)
                self.config_manager = None
        except Exception as e:
            print(f"配置管理器初始化失败: {e}")
            import traceback
            traceback.print_exc()
            # 如果config_manager不可用，使用默认配置
            self.config = self._get_default_config(config_path)
            self.config_manager = None

        # 初始化LLM
        self.llm = None
        self.model_type = None
        self._initialize_llm()

        # 配置文献获取参数
        self.download_timeout = self.config.get("download_timeout", 60)
        self.max_retries = self.config.get("max_retries", 3)
        self.retry_delay_min = self.config.get("retry_delay_min", 4)
        self.retry_delay_max = self.config.get("retry_delay_max", 10)

        # 配置分析模式
        self.analysis_mode = self.config.get("analysis_mode", "full")  # "full" or "quick" or "subjective"

        # 初始化进度跟踪
        self.total_citations = 0
        self.processed_citations = 0
        self.progress_file = "analysis_progress.json"

    def _create_doc_from_markdown(self):
        """
        从Markdown内容创建模拟的docx.Document对象
        """
        import docx
        from docx import Documen
        from docx.text.paragraph import Paragraph
        from docx.section import Section

        self.doc = Document()

        # 清空默认段落
        if self.doc.paragraphs:
            # 删除默认段落
            for paragraph in self.doc.paragraphs:
                p = paragraph._elemen
                p.getparent().remove(p)

        # 按行分割Markdown内容并添加为段落
        if self.md_content is not None:
            lines = self.md_content.split('\n')
            for line in lines:
                # 添加段落到文档
                paragraph = self.doc.add_paragraph()
                paragraph.text = line.strip()
        else:
            # 如果md_content为None，添加一个空段落或适当的提示
            paragraph = self.doc.add_paragraph()
            paragraph.text = "内容为空"

    def _get_default_config(self, config_path):
        """当config_manager不可用时，提供默认配置"""
        default_config = {
            "model": "qwen",
            "api_key": "your-api-key",
            "model_name": "qwen-plus",
            "api_url": None,
            "download_timeout": 60,
            "max_retries": 3,
            "retry_delay_min": 4,
            "retry_delay_max": 10,
            "analysis_mode": "subjective",  # 默认使用主观分析模式
            "use_advanced_extraction": True
        }

        # 尝试从文件加载配置
        if os.path.exists(config_path):
            try:
                with open(config_path, 'r', encoding='utf-8') as f:
                    file_config = json.load(f)
                    default_config.update(file_config)
            except Exception as e:
                print(f"警告: 无法加载配置文件 {config_path}: {e}")

        return default_config

    def _initialize_llm(self):
        """初始化语言模型 - 使用新的AI服务模块"""
        # 从ai模块导入AIClien
        try:
            from ai.ai_client import AIClien
            # 根据配置决定使用哪种AI服务
            if self.config["model"] == "gpt":
                if not self.config.get("api_key") or self.config.get("api_key") == "your-api-key":
                    print("警告：未设置有效的GPT_API_KEY，将跳过AI相关性分析")
                    self.ai_client = None
                else:
                    self.ai_client = AIClient(
                        provider_type="openai",
                        api_key=self.config.get("api_key"),
                        model=self.config.get("model_name", "gpt-3.5-turbo"),
                        base_url=self.config.get("api_url", "https://api.openai.com/v1")
                    )
                    self.model_type = "gpt"
                    self.model_name = self.config.get("model_name", "gpt-3.5-turbo")
            elif self.config["model"] == "qwen":
                if not self.config.get("api_key") or self.config.get("api_key") == "your-api-key":
                    print("警告：未设置有效的QWEN_API_KEY，将跳过AI相关性分析")
                    self.ai_client = None
                else:
                    self.ai_client = AIClient(
                        provider_type="dashscope",
                        api_key=self.config.get("api_key"),
                        model=self.config.get("model_name", "qwen-plus")
                    )
                    self.model_type = "qwen"
                    self.model_name = self.config.get("model_name", "qwen-plus")
        except ImportError:
            print("警告：无法导入AI服务模块，将跳过AI相关性分析")
            self.ai_client = None
        except Exception as e:
            print(f"警告：AI客户端初始化失败: {e}")
            self.ai_client = None

    def extract_citations_and_references(self):
        """提取文中的引用和参考文献列表"""
        # 区分处理Word文档和PDF文档
        if self.doc_path.lower().endswith('.pdf'):
            # 对于PDF文档，使用从Markdown内容提取的引文
            expanded_citations = set()

            # 处理数字引用（从模拟的docx文档中提取）
            full_text = []
            for paragraph in self.doc.paragraphs:
                full_text.append(paragraph.text)

            full_text_str = '\n'.join(full_text)

            # 提取文中的引用（包括单个引用和范围引用，例如[1], [2-5]等格式）
            citation_pattern = r'\[\d+(?:-\d+)?\]'
            raw_citations = re.findall(citation_pattern, full_text_str)

            # 展开范围引用为单个引用
            for citation in raw_citations:
                if '-' in citation:
                    # 处理范围引用，如[1-3]
                    match = re.match(r'\[(\d+)-(\d+)\]', citation)
                    if match:
                        start = int(match.group(1))
                        end = int(match.group(2))
                        # 展开范围引用为单个引用
                        for i in range(start, end + 1):
                            expanded_citations.add(f'[{i}]')
                else:
                    # 单个引用，直接添加
                    expanded_citations.add(citation)

            # 如果启用了AI增强引用提取，且有Markdown提取模块可用
            if (AI_ENHANCED_EXTRACTION_AVAILABLE and
                MARKDOWN_EXTRACTION_AVAILABLE and
                self.config.get("use_advanced_extraction", False)):
                try:
                    # 准备AI优化的配置
                    ai_config = {
                        "api_key": self.config.get("api_key"),
                        "model_name": self.config.get("model_name", "qwen-plus")
                    }

                    # 从Markdown内容提取作者年份格式的引用
                    if hasattr(self, 'md_content') and self.md_content:
                        markdown_citations = extract_citations_from_markdown(self.md_content, ai_config)
                        # 将提取到的引用添加到引用列表中（使用特殊标识）
                        for citation in markdown_citations:
                            formatted_citation = f"[AUTH:{citation}]"
                            expanded_citations.add(formatted_citation)
                    else:
                        print("警告：PDF文档未加载Markdown内容")

                except Exception as e:
                    print(f"PDF高级引用提取失败: {e}")
                    import traceback
                    traceback.print_exc()
        else:
            # 对于Word文档，使用原来的逻辑
            full_text = []
            for paragraph in self.doc.paragraphs:
                full_text.append(paragraph.text)

            full_text_str = '\n'.join(full_text)

            # 提取文中的引用（包括单个引用和范围引用，例如[1], [2-5]等格式）
            citation_pattern = r'\[\d+(?:-\d+)?\]'
            raw_citations = re.findall(citation_pattern, full_text_str)

            # 展开范围引用为单个引用
            expanded_citations = set()
            for citation in raw_citations:
                if '-' in citation:
                    # 处理范围引用，如[1-3]
                    match = re.match(r'\[(\d+)-(\d+)\]', citation)
                    if match:
                        start = int(match.group(1))
                        end = int(match.group(2))
                        # 展开范围引用为单个引用
                        for i in range(start, end + 1):
                            expanded_citations.add(f'[{i}]')
                else:
                    # 单个引用，直接添加
                    expanded_citations.add(citation)

            # 如果启用了AI增强引用提取，则使用新的提取方式
            if AI_ENHANCED_EXTRACTION_AVAILABLE and self.config.get("use_advanced_extraction", False):
                try:
                    # 准备AI优化的配置
                    ai_config = {
                        "api_key": self.config.get("api_key"),
                        "model_name": self.config.get("model_name", "qwen-plus")
                    }

                    # 提取作者年份格式的引用
                    author_year_citations = extract_references(self.doc_path, ai_config)
                    western_citations = extract_western_references(self.doc_path, ai_config)

                    # 将提取到的引用添加到引用列表中（使用特殊标识）
                    for citation in author_year_citations + western_citations:
                        formatted_citation = f"[AUTH:{citation}]"
                        expanded_citations.add(formatted_citation)
                except Exception as e:
                    print(f"Word文档高级引用提取失败: {e}")

        # 转换为列表并去重
        self.citations = list(expanded_citations)

        # 查找参考文献部分
        references_start = None
        references_end = None

        # 从后往前查找参考文献标题，优先选择后面的
        for i in range(len(self.doc.paragraphs) - 1, -1, -1):
            text = self.doc.paragraphs[i].text.strip()
            if '参考文献' in text or 'References' in text:
                references_start = i
                # 如果这个参考文献在文档的后半部分，就使用它
                if i > len(self.doc.paragraphs) // 2:
                    break

        if references_start is None:
            print("未找到参考文献标题")
            return

        # 查找结束标记
        for i in range(references_start + 1, len(self.doc.paragraphs)):
            text = self.doc.paragraphs[i].text.strip()
            if ('附录' in text or '致谢' in text or '作者简历' in text) and len(text) < 20:
                references_end = i
                break

        # 如果没有找到结束标记，设置一个合理的范围
        if references_end is None:
            references_end = min(len(self.doc.paragraphs), references_start + 150)

        # 提取所有可能的参考文献条目 (使用check_reference_extraction.py的宽松条件)
        for i in range(references_start + 1, references_end):
            text = self.doc.paragraphs[i].text.strip()
            if text:
                # 检查是否可能为参考文献条目
                # 使用更宽松的条件
                if (len(text) > 10 and
                    ('.' in text or '[' in text) and
                    re.search(r'\d{4}', text) and
                    not text.startswith('附录') and
                    not text.startswith('致谢') and
                    not text.startswith('作者简历') and
                    not text.startswith('图') and
                    not text.startswith('表') and
                    not re.match(r'^\d+\.?[\s]*[您的]\s*', text) and
                    '目录' not in text):
                    # 尝试从参考文献条目中提取DOI或URL
                    doi_match = re.search(r'doi:\s*([^\s,.;]+)', text, re.IGNORECASE)
                    url_match = re.search(r'https?://[^\s,.;]+', text)
                    doi = doi_match.group(1) if doi_match else None
                    url = url_match.group(0) if url_match else None

                    # 保存参考文献及其元数据
                    self.references.append({
                        'text': text,
                        'doi': doi,
                        'url': url
                    })

    def check_missing_citations(self):
        """检查文中引用是否都在参考文献中出现"""
        # 提取参考文献中的编号（例如[1], [2]等）
        ref_numbers = set()
        for ref in self.references:
            # 查找参考文献条目中的编号
            number_match = re.search(r'^\[\d+\]', ref['text'])
            if number_match:
                ref_numbers.add(number_match.group())

        # 检查每个引用是否在参考文献中出现
        for citation in self.citations:
            if citation not in ref_numbers:
                self.missing_citations.append(citation)

    def check_unused_references(self):
        """检查参考文献中未被正文引用的条目"""
        # 提取正文中的引用编号
        cited_numbers = set(self.citations)

        # 检查每个参考文献条目是否被正文引用
        unused_references = []
        for ref in self.references:
            # 查找参考文献条目中的编号
            number_match = re.search(r'^\[\d+\]', ref['text'])
            if number_match:
                ref_number = number_match.group()
                if ref_number not in cited_numbers:
                    unused_references.append(ref)

        return unused_references

    def analyze_citation_relevance(self, citation: str, context: str) -> str:
        """使用AI分析引用与上下文的相关性"""
        # 处理不同格式的引用
        if citation.startswith("[AUTH:"):
            # 作者年份格式引用
            actual_citation = citation[6:-1]  # 去掉[AUTH:]前缀和后缀
            # 使用新的映射机制找到对应的参考文献条目
            from ..checker.citation_checking.reference_mapper import map_author_year_citation_to_reference
            mapping_result = map_author_year_citation_to_reference(actual_citation, self.references)
            if mapping_result:
                reference_entry = mapping_result["reference"]
                corrected_citation = mapping_result["corrected_citation"]
                original_citation = mapping_result["original_citation"]
            else:
                reference_entry = None
                corrected_citation = actual_citation
                original_citation = actual_citation
        else:
            # 数字格式引用
            actual_citation = citation
            corrected_citation = actual_citation
            original_citation = actual_citation
            # 找到对应的参考文献条目
            reference_entry = None
            for ref in self.references:
                if ref['text'].startswith(citation):
                    reference_entry = ref
                    break

        if not reference_entry:
            return "未找到对应的参考文献条目"

        # 检查AI客户端是否可用
        if not self.ai_client:
            return "无法进行AI分析：未配置有效的AI客户端"

        # 尝试获取论文信息（标题和摘要）
        paper_info = self.fetch_paper_info(reference_entry)

        # 根据分析模式决定是否进行全文分析
        full_text_content = ""
        if self.analysis_mode == "full":
            # 在完整模式下，尝试获取论文全文内容
            full_text_content = self.fetch_paper_content(reference_entry)

        # 创建提示模板
        if paper_info and ('title' in paper_info or 'abstract' in paper_info):
            # 如果获取到了论文标题或摘要，使用这些信息
            title = paper_info.get('title', '无')
            abstract = paper_info.get('abstract', '无')

            # 在完整模式下，如果有全文内容则添加到提示中
            if self.analysis_mode == "full" and full_text_content:
                prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

参考文献论文信息：
标题：{title}
摘要：{abstract}

论文全文内容（前几页）：
{full_text_content[:2000]}...（内容截断）

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""
            else:
                prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

参考文献论文信息：
标题：{title}
摘要：{abstract}

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""
        else:
            # 如果没有获取到论文信息，使用基本提示
            prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""

            # 如果有DOI或URL，提示用户可以尝试下载全文
            if reference_entry.get('doi') or reference_entry.get('url'):
                prompt_text += f"""

注意：该参考文献有DOI或URL，可以尝试下载全文进行更准确的分析。
DOI: {reference_entry.get('doi', '无')}
URL: {reference_entry.get('url', '无')}"""

        # 限制提示文本长度，避免超过模型的最大输入长度
        max_prompt_length = 3000  # 根据模型的最大输入长度调整
        if len(prompt_text) > max_prompt_length:
            prompt_text = prompt_text[:max_prompt_length] + "\n\n...（内容已截断）"

        try:
            response = self.ai_client.generate(prompt_text, max_tokens=500, temperature=0)
            return response
        except Exception as e:
            return f"AI分析时出错：{str(e)}"

    @retry_strategy
    def fetch_paper_info(self, reference_entry: Dict) -> Dict:
        """
        尝试根据DOI或标题获取论文的标题和摘要信息
        """
        try:
            # 生成缓存文件名
            cache_key = hashlib.md5((str(reference_entry) + "_info").encode('utf-8')).hexdigest()
            cache_file = os.path.join(CACHE_DIR, f"{cache_key}.json")

            # 检查缓存是否存在
            if os.path.exists(cache_file):
                logger.info(f"从缓存中读取信息: {cache_file}")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return json.load(f)

            # 初始化Semantic Scholar客户端
            import semanticscholar as sch
            s2 = sch.SemanticScholar(timeout=self.download_timeout)

            # 获取论文信息
            paper_info = None

            # 首先尝试通过Semantic Scholar API获取论文信息
            if reference_entry.get('doi'):
                try:
                    logger.info(f"尝试通过Semantic Scholar获取DOI {reference_entry['doi']} 的信息")
                    paper_info = s2.get_paper(f"DOI:{reference_entry['doi']}")
                    logger.info(f"通过DOI {reference_entry['doi']} 获取Semantic Scholar信息成功")
                except requests.exceptions.RequestException as e:
                    if self._handle_rate_limit(e.response):
                        raise e  # 重新抛出异常以触发重试
                    logger.warning(f"通过DOI获取Semantic Scholar信息时发生网络错误: {e}")
                except Exception as e:
                    logger.warning(f"通过DOI获取Semantic Scholar信息失败: {e}")

            # 如果没有通过DOI获取到信息，尝试通过Semantic Scholar的标题搜索
            if not paper_info:
                # 提取参考文献标题（简单地取前几个单词）
                title_match = re.search(r'^\[\d+\]\s*(.+?)(?:\.\s|\.\s*$)', reference_entry['text'])
                if title_match:
                    title = title_match.group(1)
                    try:
                        logger.info(f"尝试通过Semantic Scholar搜索标题 '{title}'")
                        # 使用标题搜索论文
                        results = s2.search_paper(title, limit=1)
                        if results and len(results) > 0:
                            paper_info = results[0]
                            logger.info(f"通过标题 '{title}' 搜索Semantic Scholar信息成功")
                    except requests.exceptions.RequestException as e:
                        if self._handle_rate_limit(e.response):
                            raise e  # 重新抛出异常以触发重试
                        logger.warning(f"通过标题搜索Semantic Scholar信息时发生网络错误: {e}")
                    except Exception as e:
                        logger.warning(f"通过标题搜索Semantic Scholar信息失败: {e}")

            # 如果获取到了信息，提取标题和摘要
            if paper_info:
                info = {}
                if hasattr(paper_info, 'title') and paper_info.title:
                    info['title'] = paper_info.title
                if hasattr(paper_info, 'abstract') and paper_info.abstract:
                    info['abstract'] = paper_info.abstrac

                # 保存到缓存
                with open(cache_file, 'w', encoding='utf-8') as f:
                    json.dump(info, f, ensure_ascii=False, indent=2)
                logger.info(f"信息已缓存到: {cache_file}")

                return info

            logger.warning("无法获取论文信息")
            return {}
        except requests.exceptions.RequestException as e:
            logger.error(f"获取论文信息时发生网络错误: {e}")
            return {}
        except Exception as e:
            logger.error(f"获取论文信息时出错: {e}")
            return {}

    @retry_strategy
    def fetch_paper_content(self, reference_entry: Dict) -> str:
        """
        尝试根据DOI或URL下载论文并提取文本内容
        """
        try:
            # 生成缓存文件名
            cache_key = hashlib.md5(str(reference_entry).encode('utf-8')).hexdigest()
            cache_file = os.path.join(CACHE_DIR, f"{cache_key}.txt")

            # 检查缓存是否存在
            if os.path.exists(cache_file):
                logger.info(f"从缓存中读取内容: {cache_file}")
                with open(cache_file, 'r', encoding='utf-8') as f:
                    return f.read()

            # 初始化Semantic Scholar客户端
            import semanticscholar as sch
            s2 = sch.SemanticScholar(timeout=self.download_timeout)

            # 获取论文的URL
            paper_url = None

            # 首先尝试通过Semantic Scholar API获取论文信息
            paper_info = None
            if reference_entry.get('doi'):
                try:
                    logger.info(f"尝试通过Semantic Scholar获取DOI {reference_entry['doi']} 的信息")
                    paper_info = s2.get_paper(f"DOI:{reference_entry['doi']}")
                    logger.info(f"通过DOI {reference_entry['doi']} 获取Semantic Scholar信息成功")
                except requests.exceptions.RequestException as e:
                    if self._handle_rate_limit(e.response):
                        raise e  # 重新抛出异常以触发重试
                    logger.warning(f"通过DOI获取Semantic Scholar信息时发生网络错误: {e}")
                except Exception as e:
                    logger.warning(f"通过DOI获取Semantic Scholar信息失败: {e}")

            # 如果没有通过DOI获取到信息，尝试通过Crossref API获取论文信息
            if not paper_info and reference_entry.get('doi'):
                try:
                    logger.info(f"尝试通过Crossref获取DOI {reference_entry['doi']} 的信息")
                    # 使用Crossref API获取论文信息
                    import crossref
                    work = crossref.restful.Works()
                    result = work.doi(reference_entry['doi'])
                    if result and 'URL' in result:
                        paper_url = result['URL']
                        logger.info(f"通过DOI {reference_entry['doi']} 获取Crossref信息成功，URL: {paper_url}")
                except requests.exceptions.RequestException as e:
                    if self._handle_rate_limit(e.response):
                        raise e  # 重新抛出异常以触发重试
                    logger.warning(f"通过DOI获取Crossref信息时发生网络错误: {e}")
                except Exception as e:
                    logger.warning(f"通过DOI获取Crossref信息失败: {e}")

            # 如果没有通过DOI获取到信息，尝试通过Semantic Scholar的标题搜索
            if not paper_info:
                # 提取参考文献标题（简单地取前几个单词）
                title_match = re.search(r'^\[\d+\]\s*(.+?)(?:\.\s|\.\s*$)', reference_entry['text'])
                if title_match:
                    title = title_match.group(1)
                    try:
                        logger.info(f"尝试通过Semantic Scholar搜索标题 '{title}'")
                        # 使用标题搜索论文
                        results = s2.search_paper(title, limit=1)
                        if results and len(results) > 0:
                            paper_info = results[0]
                            logger.info(f"通过标题 '{title}' 搜索Semantic Scholar信息成功")
                    except requests.exceptions.RequestException as e:
                        if self._handle_rate_limit(e.response):
                            raise e  # 重新抛出异常以触发重试
                        logger.warning(f"通过标题搜索Semantic Scholar信息时发生网络错误: {e}")
                    except Exception as e:
                        logger.warning(f"通过标题搜索Semantic Scholar信息失败: {e}")

            # 如果通过Semantic Scholar获取到了信息，尝试获取PDF URL
            if paper_info and hasattr(paper_info, 'openAccessPdf') and paper_info.openAccessPdf:
                paper_url = paper_info.openAccessPdf.get('url')
                logger.info(f"从Semantic Scholar获取PDF URL: {paper_url}")

            # 如果没有通过Semantic Scholar获取到PDF URL，尝试使用DOI或URL
            if not paper_url:
                if reference_entry.get('doi'):
                    # 尝试通过DOI解析URL
                    paper_url = f"https://doi.org/{reference_entry['doi']}"
                    logger.info(f"使用DOI构造URL: {paper_url}")
                elif reference_entry.get('url'):
                    paper_url = reference_entry['url']
                    logger.info(f"使用参考文献中的URL: {paper_url}")

            if not paper_url:
                logger.warning("无法获取论文URL")
                return ""

            # 发送HTTP请求获取论文内容
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }

            # 如果是DOI，可能需要重定向到实际的PDF链接
            if "doi.org" in paper_url:
                logger.info(f"发送HEAD请求到DOI链接: {paper_url}")
                response = requests.head(paper_url, headers=headers, allow_redirects=True, timeout=self.download_timeout)
                # 检查最终重定向的URL是否是PDF
                if response.headers.get('content-type', '').startswith('application/pdf'):
                    pdf_url = response.url
                    logger.info(f"通过DOI重定向获取PDF URL: {pdf_url}")
                else:
                    # 如果不是PDF，可能需要通过其他方式获取PDF链接
                    # 这里我们简单地尝试直接访问DOI链接
                    pdf_url = paper_url
                    logger.info(f"DOI链接未重定向到PDF，直接使用: {pdf_url}")
            else:
                pdf_url = paper_url

            # 下载PDF文件
            logger.info(f"开始下载PDF: {pdf_url}")
            response = requests.get(pdf_url, headers=headers, timeout=self.download_timeout)
            if response.status_code == 200 and response.headers.get('content-type', '').startswith('application/pdf'):
                # 保存PDF到临时文件
                temp_pdf_path = f"temp_{reference_entry['text'][:10].replace('[', '').replace(']', '').replace(' ', '_')}.pdf"
                logger.info(f"PDF下载成功，保存到: {temp_pdf_path}")
                with open(temp_pdf_path, 'wb') as f:
                    f.write(response.content)

                # 使用pymupdf4llm解析PDF
                try:
                    logger.info(f"开始解析PDF: {temp_pdf_path}")
                    # 提取前几页的文本内容（避免处理过长的文档）
                    pages = pymupdf4llm.to_markdown(temp_pdf_path, pages=[1, 2, 3, 4, 5])
                    # 简单地将页面内容连接起来
                    content = "\n".join(pages) if isinstance(pages, list) else str(pages)
                    logger.info("PDF解析成功")

                    # 清理临时文件
                    os.remove(temp_pdf_path)
                    logger.info(f"临时文件 {temp_pdf_path} 已清理")

                    # 保存到缓存
                    with open(cache_file, 'w', encoding='utf-8') as f:
                        f.write(content)
                    logger.info(f"内容已缓存到: {cache_file}")

                    return conten
                except Exception as e:
                    logger.error(f"解析PDF时出错: {e}")
                    # 清理临时文件
                    if os.path.exists(temp_pdf_path):
                        os.remove(temp_pdf_path)
                        logger.info(f"临时文件 {temp_pdf_path} 已清理")
                    return ""
            else:
                logger.warning(f"无法下载PDF: {pdf_url}, 状态码: {response.status_code}, 内容类型: {response.headers.get('content-type', '')}")
                return ""
        except requests.exceptions.RequestException as e:
            logger.error(f"下载论文时发生网络错误: {e}")
            return ""
        except Exception as e:
            logger.error(f"下载或解析论文时出错: {e}")
            return ""

    def find_context_around_citation(self, citation: str) -> str:
        """查找引用周围的上下文（基于完整段落）"""
        # 查找参考文献部分的开始位置
        references_start_idx = None
        for i, paragraph in enumerate(self.doc.paragraphs):
            if '参考文献' in paragraph.text or 'References' in paragraph.text:
                references_start_idx = i
                break

        # 首先在正文中查找包含该引用的段落或表格
        # 但要注意，如果原文中使用的是范围引用如[1-3]，那么单独的[1]、[2]、[3]可能不存在
        if references_start_idx is not None:
            # 在正文中查找所有包含该引用的段落
            body_paragraphs_with_citation = []
            for i, paragraph in enumerate(self.doc.paragraphs[:references_start_idx]):
                # 检查段落是否包含该引用（直接引用或范围引用）
                if (citation in paragraph.text) and paragraph.text.strip():
                    body_paragraphs_with_citation.append((i, paragraph))
                # 针对展开的引用，检查是否存在对应的范围引用
                elif self._is_in_range_citation(citation, paragraph.text):
                    body_paragraphs_with_citation.append((i, paragraph))

            # 在表格中查找包含该引用的单元格
            table_contexts = []
            if hasattr(self.doc, 'tables'):
                for table in self.doc.tables:
                    # 检查表格是否在正文部分（在参考文献之前）
                    # 这里我们简化处理，假设表格都在正文部分
                    for row in table.rows:
                        for cell in row.cells:
                            if (citation in cell.text) and cell.text.strip():
                                table_contexts.append(cell.text.strip())
                            # 针对展开的引用，检查是否存在对应的范围引用
                            elif self._is_in_range_citation(citation, cell.text):
                                table_contexts.append(cell.text.strip())

            # 如果在正文中找到了包含该引用的段落
            if body_paragraphs_with_citation:
                # 选择第一个包含该引用的正文段落
                i, paragraph = body_paragraphs_with_citation[0]

                # 提取当前段落以及前后段落作为上下文
                context_paragraphs = []

                # 添加前一个段落（如果存在且不为空）
                if i > 0 and self.doc.paragraphs[i-1].text.strip():
                    context_paragraphs.append(self.doc.paragraphs[i-1].text.strip())

                # 添加当前段落（包含引用标记的段落）
                context_paragraphs.append(paragraph.text.strip())

                # 添加后一个段落（如果存在且不为空）
                if i < references_start_idx - 1 and self.doc.paragraphs[i+1].text.strip():
                    context_paragraphs.append(self.doc.paragraphs[i+1].text.strip())

                # 将段落连接成完整的上下文
                context = " ".join(context_paragraphs)
                return contex

            # 如果在表格中找到了包含该引用的内容
            elif table_contexts:
                # 返回表格单元格的内容作为上下文
                return table_contexts[0]

        # 如果在正文中没有找到，不要去参考文献部分查找单个引用条目
        # 因为参考文献条目不是上下文，而是引用的目标

        # 回退到字符级别的方法
        full_text = []
        for paragraph in self.doc.paragraphs:
            full_text.append(paragraph.text)

        # 也要包含表格中的内容
        if hasattr(self.doc, 'tables'):
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)

        full_text_str = '\n'.join(full_text)

        # 查找引用位置
        citation_pos = full_text_str.find(citation)
        if citation_pos == -1:
            return ""

        # 提取引用前后各100个字符作为上下文
        start = max(0, citation_pos - 100)
        end = min(len(full_text_str), citation_pos + len(citation) + 100)
        context = full_text_str[start:end]

        return contex

    def _is_in_range_citation(self, citation: str, text: str) -> bool:
        """检查给定的单个引用是否在范围引用中"""
        # 匹配范围引用模式，如[1-3], [10-15]等
        range_pattern = r'\[(\d+)-(\d+)\]'
        matches = re.finditer(range_pattern, text)

        # 提取引用编号
        citation_number = int(re.search(r'\[(\d+)\]', citation).group(1))

        for match in matches:
            start = int(match.group(1))
            end = int(match.group(2))
            # 检查引用编号是否在范围中
            if start <= citation_number <= end:
                return True

        return False

    def _analyze_single_citation(self, citation: str) -> tuple:
        """分析单个引用的相关性，用于多线程处理"""
        start_time = time.time()

        # 计算当前进度百分比
        if self.total_citations > 0:
            progress_percentage = int((self.processed_citations / self.total_citations) * 100)
            print(f"开始分析引用 {citation} 的相关性... ({self.processed_citations}/{self.total_citations}, {progress_percentage}%)")
        else:
            print(f"开始分析引用 {citation} 的相关性...")
        sys.stdout.flush()

        # 获取上下文信息
        context = self.find_context_around_citation(citation)

        if citation not in self.missing_citations:
            # 根据分析模式选择合适的分析函数
            if self.analysis_mode == "quick":
                analysis = self.analyze_citation_relevance_quick(citation, context)
            elif self.analysis_mode == "subjective":
                analysis = self.analyze_citation_relevance_subjective(citation, context)
            else:
                analysis = self.analyze_citation_relevance(citation, context)
            end_time = time.time()

            # 计算当前进度百分比
            if self.total_citations > 0:
                progress_percentage = int((self.processed_citations / self.total_citations) * 100)
                print(f"完成分析引用 {citation} 的相关性，耗时 {end_time - start_time:.2f} 秒 ({self.processed_citations}/{self.total_citations}, {progress_percentage}%)")
            else:
                print(f"完成分析引用 {citation} 的相关性，耗时 {end_time - start_time:.2f} 秒")
            sys.stdout.flush()

            # 更新进度
            self.processed_citations += 1
            self._update_progress()

            return (citation, analysis, context)

        end_time = time.time()

        # 计算当前进度百分比
        if self.total_citations > 0:
            progress_percentage = int((self.processed_citations / self.total_citations) * 100)
            print(f"完成分析引用 {citation} 的相关性，耗时 {end_time - start_time:.2f} 秒 ({self.processed_citations}/{self.total_citations}, {progress_percentage}%)")
        else:
            print(f"完成分析引用 {citation} 的相关性，耗时 {end_time - start_time:.2f} 秒")
        sys.stdout.flush()

        # 更新进度
        self.processed_citations += 1
        self._update_progress()

        return (citation, "引用未在参考文献中找到", context)

    def _update_progress(self):
        """更新分析进度"""
        if self.total_citations > 0:
            progress = {
                "processed": self.processed_citations,
                "total": self.total_citations,
                "percentage": int((self.processed_citations / self.total_citations) * 100)
            }

            # 将进度信息写入文件
            with open(self.progress_file, 'w', encoding='utf-8') as f:
                json.dump(progress, f)

    def analyze_citation_relevance_quick(self, citation: str, context: str) -> str:
        """快速分析引用与上下文的相关性（仅使用标题和摘要）"""
        # 处理不同格式的引用
        if citation.startswith("[AUTH:"):
            # 作者年份格式引用
            actual_citation = citation[6:-1]  # 去掉[AUTH:]前缀和后缀
            # 使用新的映射机制找到对应的参考文献条目
            from ..checker.citation_checking.reference_mapper import map_author_year_citation_to_reference
            mapping_result = map_author_year_citation_to_reference(actual_citation, self.references)
            if mapping_result:
                reference_entry = mapping_result["reference"]
                corrected_citation = mapping_result["corrected_citation"]
                original_citation = mapping_result["original_citation"]
            else:
                reference_entry = None
                corrected_citation = actual_citation
                original_citation = actual_citation
        else:
            # 数字格式引用
            actual_citation = citation
            corrected_citation = actual_citation
            original_citation = actual_citation
            # 找到对应的参考文献条目
            reference_entry = None
            for ref in self.references:
                if ref['text'].startswith(citation):
                    reference_entry = ref
                    break

        if not reference_entry:
            return "未找到对应的参考文献条目"

        # 如果没有配置模型，返回默认信息
        if not self.model_type:
            return "无法进行AI分析：未配置有效的API密钥"

        # 尝试获取论文信息（标题和摘要）
        paper_info = self.fetch_paper_info(reference_entry)

        # 创建提示模板
        if paper_info and ('title' in paper_info or 'abstract' in paper_info):
            # 如果获取到了论文标题或摘要，使用这些信息
            title = paper_info.get('title', '无')
            abstract = paper_info.get('abstract', '无')

            prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

参考文献论文信息：
标题：{title}
摘要：{abstract}

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""
        else:
            # 如果没有获取到论文信息，使用基本提示
            prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""

        # 限制提示文本长度，避免超过模型的最大输入长度
        max_prompt_length = 3000  # 根据模型的最大输入长度调整
        if len(prompt_text) > max_prompt_length:
            prompt_text = prompt_text[:max_prompt_length] + "\n\n...（内容已截断）"

        try:
            if self.ai_client:
                response = self.ai_client.generate(prompt_text, max_tokens=500, temperature=0)
                return response
            else:
                return "无法进行AI分析：未配置有效的AI客户端"
        except Exception as e:
            return f"AI分析时出错：{str(e)}"

    def analyze_citation_relevance_subjective(self, citation: str, context: str) -> str:
        """主观分析引用与上下文的相关性（完全依赖AI判断，不使用外部API）"""
        # 处理不同格式的引用
        if citation.startswith("[AUTH:"):
            # 作者年份格式引用
            actual_citation = citation[6:-1]  # 去掉[AUTH:]前缀和后缀
            # 使用新的映射机制找到对应的参考文献条目
            from ..checker.citation_checking.reference_mapper import map_author_year_citation_to_reference
            mapping_result = map_author_year_citation_to_reference(actual_citation, self.references)
            if mapping_result:
                reference_entry = mapping_result["reference"]
                corrected_citation = mapping_result["corrected_citation"]
                original_citation = mapping_result["original_citation"]
            else:
                reference_entry = None
                corrected_citation = actual_citation
                original_citation = actual_citation
        else:
            # 数字格式引用
            actual_citation = citation
            corrected_citation = actual_citation
            original_citation = actual_citation
            # 找到对应的参考文献条目
            reference_entry = None
            for ref in self.references:
                if ref['text'].startswith(citation):
                    reference_entry = ref
                    break

        if not reference_entry:
            return "未找到对应的参考文献条目"

        # 如果没有配置模型，返回默认信息
        if not self.model_type:
            return "无法进行AI分析：未配置有效的API密钥"

        # 创建提示模板，不使用任何外部API获取的信息
        prompt_text = f"""
请分析以下学术论文中的引用是否与上下文相关：

引用编号：{actual_citation}
参考文献条目：{reference_entry['text']}
引用上下文：{context}

请严格按照以下格式回答：
1. 相关性判断：相关 / 不相关
2. 分析理由：详细说明引用与上下文的相关性分析理由，包括引用文献的主题内容、上下文讨论的内容、两者之间的关联性或不匹配之处
3. 问题说明：如果相关则写"无"，如果不相关则指出可能的正确引用或说明具体问题

请以简洁明了的方式回答，严格按照上述格式输出。
"""

        # 限制提示文本长度，避免超过模型的最大输入长度
        max_prompt_length = 3000  # 根据模型的最大输入长度调整
        if len(prompt_text) > max_prompt_length:
            prompt_text = prompt_text[:max_prompt_length] + "\n\n...（内容已截断）"

        try:
            if self.ai_client:
                response = self.ai_client.generate(prompt_text, max_tokens=500, temperature=0)
                return response
            else:
                return "无法进行AI分析：未配置有效的AI客户端"
        except Exception as e:
            return f"AI分析时出错：{str(e)}"

    def _handle_rate_limit(self, response):
        """处理API速率限制"""
        if response.status_code == 429:
            # 从响应头获取重试时间，如果没有则随机等待
            retry_after = response.headers.get('Retry-After')
            if retry_after:
                wait_time = int(retry_after)
            else:
                wait_time = random.randint(5, 15)
            logger.warning(f"遇到API速率限制 (429)，等待 {wait_time} 秒后重试...")
            time.sleep(wait_time)
            return True
        return False

    def generate_report(self):
        """生成合规性检查报告"""
        self.extract_citations_and_references()
        self.check_missing_citations()

        # 检查未被引用的参考文献条目
        unused_references = self.check_unused_references()

        # 初始化进度跟踪
        self.total_citations = len(self.citations)
        self.processed_citations = 0

        # 清除之前的进度文件
        if os.path.exists(self.progress_file):
            os.remove(self.progress_file)

        report = []
        report.append("<h1>论文引用合规性检查报告</h1>\n")

        # 分类显示未出现的引用
        missing_numeric = [c for c in self.missing_citations if not c.startswith("[AUTH:")]
        missing_author_year = [c for c in self.missing_citations if c.startswith("[AUTH:")]

        # 报告未出现的数字引用
        if missing_numeric:
            report.append("<h2>未在参考文献中出现的数字引用</h2>\n<ul>\n")
            for citation in missing_numeric:
                report.append(f"<li>{citation}</li>\n")
            report.append("</ul>\n")
        else:
            report.append("<h2>未在参考文献中出现的数字引用</h2>\n")
            report.append("<p>所有文中的数字引用都在参考文献中找到了对应条目。</p>\n")

        # 报告未出现的作者年份引用
        if missing_author_year:
            report.append("<h2>未在参考文献中出现的作者年份引用</h2>\n<ul>\n")
            for citation in missing_author_year:
                # 显示时去掉[AUTH:]前缀
                display_citation = citation[6:-1] if citation.startswith("[AUTH:") else citation
                report.append(f"<li>{display_citation}</li>\n")
            report.append("</ul>\n")
        else:
            report.append("<h2>未在参考文献中出现的作者年份引用</h2>\n")
            report.append("<p>所有文中的作者年份引用都在参考文献中找到了对应条目。</p>\n")

        # 报告未被引用的参考文献
        if unused_references:
            report.append("<h2>参考文献中未被正文引用的条目</h2>\n<ul>\n")
            for ref in unused_references:
                # 提取参考文献编号
                number_match = re.search(r'^\[\d+\]', ref['text'])
                ref_number = number_match.group() if number_match else "未知"
                report.append(f"<li>{ref_number}: {ref['text']}</li>\n")
            report.append("</ul>\n")
        else:
            report.append("<h2>参考文献中未被正文引用的条目</h2>\n")
            report.append("<p>所有参考文献条目都被正文引用。</p>\n")

        # 分析引用相关性 - 使用多线程
        report.append("<h2>引用相关性分析</h2>\n")

        # 确定线程数，最多不超过引用数量
        num_threads = min(10, len(self.citations))  # 最多使用10个线程

        if num_threads > 0 and self.llm:
            print(f"开始使用 {num_threads} 个线程进行引用相关性分析...")
            with ThreadPoolExecutor(max_workers=num_threads) as executor:
                # 提交所有任务
                future_to_citation = {
                    executor.submit(self._analyze_single_citation, citation): citation
                    for citation in self.citations
                }

                # 收集结果
                analysis_results = {}
                context_results = {}
                for future in as_completed(future_to_citation):
                    result = future.result()
                    if len(result) == 3:
                        citation, analysis, context = resul
                    else:
                        citation, analysis = resul
                        context = ""
                    analysis_results[citation] = analysis
                    context_results[citation] = contex

                # 按引用编号排序结果
                def extract_citation_number(citation):
                    # 如果是作者年份格式的引用，返回一个大的数字使其排在最后
                    if citation.startswith("[AUTH:"):
                        return 999999
                    # 否则提取数字引用编号
                    match = re.search(r'\[(\d+)\]', citation)
                    return int(match.group(1)) if match else 999999

                sorted_citations = sorted(self.citations, key=extract_citation_number)

                # 添加到报告
                for citation in sorted_citations:
                    if citation not in self.missing_citations:
                        analysis = analysis_results.get(citation, "分析失败")
                        context = context_results.get(citation, "")
                        # 转义HTML特殊字符
                        context_escaped = context.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        analysis_escaped = analysis.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                        report.append(f"<h3>引用 {citation} 相关性分析</h3>\n")
                        report.append(f"<div class='context'><strong>上下文</strong>: {context_escaped}</div>\n")
                        report.append(f"<div class='analysis'>{analysis_escaped}</div>\n<br/>\n")
        else:
            # 如果没有LLM或引用数量为0，顺序执行（主要用于兼容性）
            print("未使用多线程进行分析（可能因为缺少API密钥或引用数量为0）")
            context_results = {}
            for citation in self.citations:
                if citation not in self.missing_citations:
                    context = self.find_context_around_citation(citation)
                    context_results[citation] = contex
                    analysis = self.analyze_citation_relevance(citation, context)
                    # 转义HTML特殊字符
                    context_escaped = context.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    analysis_escaped = analysis.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    report.append(f"<h3>引用 {citation} 相关性分析</h3>\n")
                    report.append(f"<div class='context'><strong>上下文</strong>: {context_escaped}</div>\n")
                    report.append(f"<div class='analysis'>{analysis_escaped}</div>\n<br/>\n")

                    # 更新进度
                    self.processed_citations += 1
                    self._update_progress()

                    # 输出进度信息
                    if self.total_citations > 0:
                        progress_percentage = int((self.processed_citations / self.total_citations) * 100)
                        print(f"完成分析引用 {citation} 的相关性 ({self.processed_citations}/{self.total_citations}, {progress_percentage}%)")
                        sys.stdout.flush()

        # 分析完成后删除进度文件
        if os.path.exists(self.progress_file):
            os.remove(self.progress_file)

        return ''.join(report)

    def cleanup(self):
        """
        清理临时文件
        """
        if self.temp_md_file and os.path.exists(self.temp_md_file):
            try:
                os.remove(self.temp_md_file)
                # 如果临时文件在临时目录中，且目录为空，也删除目录
                temp_dir = os.path.dirname(self.temp_md_file)
                if temp_dir.startswith(tempfile.gettempdir()) and os.path.isdir(temp_dir):
                    try:
                        os.rmdir(temp_dir)  # 尝试删除空目录
                    except OSError:
                        pass  # 目录不为空，不删除
            except Exception as e:
                print(f"清理临时文件失败: {e}")