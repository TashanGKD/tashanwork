"""
AI增强的文档内容提取器
包含从文档中提取引用、参考文献等的AI增强功能
"""

import re
from docx import Documen
from typing import List, Dict, Any, Optional


def extract_references_with_context(docx_path: str, config: Optional[Dict] = None) -> List[Dict[str, Any]]:
    """提取文档中的作者年份格式引用及其上下文"""
    # 多种引用格式的模式
    patterns = [
        # 格式1: 作者（年份） - 中文括号
        (r'([A-Za-z\u4e00-\u9fa5&＆\s\.]+?)\s*[（(](\d{4})[)）]', 0),
        # 格式2: (作者, 年份) - 英文括号
        (r'[（(]([A-Za-z\u4e00-\u9fa5&＆\s\.]+?)\s*[,，]?\s*(\d{4})[)）]', 1),
        # 格式3: 作者 et al.（年份）
        (r'([A-Za-z\u4e00-\u9fa5\s]+?)\s+et al\.\s*[（(](\d{4})[)）]', 0),
        # 格式4: 作者等（年份）
        (r'([A-Za-z\u4e00-\u9fa5\s]+?)\s+等\s*[（(](\d{4})[)）]', 0),
        # 格式5: 作者, 首字母. (年份) - 英文格式
        (r'([A-Z][a-z]+,\s*[A-Z]\.)\s*[（(](\d{4})[)）]', 0),
        # 格式6: 作者, 首字母. (年份) - 可能有多个作者
        (r'([A-Z][a-z]+,\s*[A-Z]\.(?:\s*&\s*[A-Z][a-z]+,\s*[A-Z]\.)*)\s*[（(](\d{4})[)）]', 0),
        # 格式7: 作者 (年份) - 英文名+空格+括号
        (r'([A-Z][a-z]+\s+[A-Z][a-z]+)\s*[（(](\d{4})[)）]', 0)
    ]

    doc = Document(docx_path)
    references = []

    # 遍历所有段落，查找引用及其上下文
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if not para_text:
            continue

        for pattern, group_type in patterns:
            matches = list(re.finditer(pattern, para_text))
            for match in matches:
                # 处理不同模式的匹配组
                if group_type == 0:  # 作者在前
                    author, year = match.group(1), match.group(2)
                else:  # 作者在括号内
                    author, year = match.group(1), match.group(2)

                # 清理作者名称
                author = author.strip()
                author = re.sub(r'^\s*[,，]\s*', '', author)
                author = re.sub(r'\s+', ' ', author)
                author = re.sub(r'[，,]\s*$', '', author)

                # 验证年份
                if not year.isdigit() or len(year) != 4:
                    continue

                # 验证作者
                if len(author) >= 2 and not author.isdigit():
                    reference = f"{author}（{year}）"

                    # 提取上下文（当前段落及前后段落）
                    context = extract_context_around_position(doc, para_idx, match.start(), match.end())

                    references.append({
                        'citation': reference,
                        'context': contex
                    })

    # 去重，保留上下文信息
    unique_references = []
    seen = set()
    for ref in references:
        if ref['citation'] not in seen:
            seen.add(ref['citation'])
            unique_references.append(ref)

    # 提取引用文本用于AI优化
    citation_texts = [ref['citation'] for ref in unique_references]

    # 使用AI优化引用列表（如果可用）
    try:
        from .ai_optimizer import optimize_citations_with_ai
        optimized_citations = optimize_citations_with_ai(citation_texts, config)

        # 更新优化后的引用文本，保持上下文信息
        optimized_references = []
        for ref in unique_references:
            original_citation = ref['citation']
            # 查找优化后的对应引用
            optimized_citation = None
            # 尝试精确匹配
            if original_citation in optimized_citations:
                optimized_citation = original_citation
            else:
                # 如果没有精确匹配，尝试模糊匹配（检查优化后的引用是否与原始引用的作者部分相关）
                original_author = original_citation.split('（')[0] if '（' in original_citation else original_citation
                for opt_citation in optimized_citations:
                    opt_author = opt_citation.split('（')[0] if '（' in opt_citation else opt_citation
                    # 检查作者部分是否相似（去除前缀后）
                    if original_author.endswith(opt_author) or opt_author in original_author:
                        optimized_citation = opt_citation
                        break

            if optimized_citation:
                optimized_references.append({
                    'citation': optimized_citation,
                    'context': ref['context']
                })
            else:
                optimized_references.append(ref)  # 如果没有优化版本，使用原始引用

        return optimized_references
    except ImportError:
        # 如果AI优化模块不可用，返回原始提取结果
        return unique_references


def extract_western_references_with_context(docx_path: str, config: Optional[Dict] = None) -> List[Dict[str, Any]]:
    """专门提取西式格式的引用及其上下文"""
    # 西式引用格式
    western_patterns = [
        # 格式: Lastname, F. (year)
        (r'([A-Z][a-z]+,\s*[A-Z]\.)\s*\((\d{4})\)', 0),
        # 格式: Lastname, F. & Lastname, G. (year)
        (r'([A-Z][a-z]+,\s*[A-Z]\.(?:\s*&\s*[A-Z][a-z]+,\s*[A-Z]\.)*)\s*\((\d{4})\)', 0),
        # 格式: Lastname (year)
        (r'([A-Z][a-z]+)\s*\((\d{4})\)', 0),
        # 格式: Lastname and Lastname (year)
        (r'([A-Z][a-z]+\s+and\s+[A-Z][a-z]+)\s*\((\d{4})\)', 0)
    ]

    doc = Document(docx_path)
    western_refs = []

    # 遍历所有段落，查找引用及其上下文
    for para_idx, paragraph in enumerate(doc.paragraphs):
        para_text = paragraph.text.strip()
        if not para_text:
            continue

        for pattern, group_type in western_patterns:
            matches = list(re.finditer(pattern, para_text))
            for match in matches:
                author, year = match.group(1), match.group(2)
                author = author.strip()
                reference = f"{author}（{year}）"

                # 提取上下文（当前段落及前后段落）
                context = extract_context_around_position(doc, para_idx, match.start(), match.end())

                western_refs.append({
                    'citation': reference,
                    'context': contex
                })

    # 提取引用文本用于AI优化
    citation_texts = [ref['citation'] for ref in western_refs]

    # 使用AI优化引用列表（如果可用）
    try:
        from .ai_optimizer import optimize_citations_with_ai
        optimized_citations = optimize_citations_with_ai(citation_texts, config)

        # 更新优化后的引用文本，保持上下文信息
        optimized_references = []
        for ref in western_refs:
            original_citation = ref['citation']
            # 查找优化后的对应引用
            optimized_citation = None
            # 尝试精确匹配
            if original_citation in optimized_citations:
                optimized_citation = original_citation
            else:
                # 如果没有精确匹配，尝试模糊匹配（检查优化后的引用是否与原始引用的作者部分相关）
                original_author = original_citation.split('（')[0] if '（' in original_citation else original_citation
                for opt_citation in optimized_citations:
                    opt_author = opt_citation.split('（')[0] if '（' in opt_citation else opt_citation
                    # 检查作者部分是否相似（去除前缀后）
                    if original_author.endswith(opt_author) or opt_author in original_author:
                        optimized_citation = opt_citation
                        break

            if optimized_citation:
                optimized_references.append({
                    'citation': optimized_citation,
                    'context': ref['context']
                })
            else:
                optimized_references.append(ref)  # 如果没有优化版本，使用原始引用

        return optimized_references
    except ImportError:
        # 如果AI优化模块不可用，返回原始提取结果
        return western_refs


def extract_context_around_position(doc: Document, para_idx: int, start_pos: int, end_pos: int, context_length: int = 100) -> str:
    """提取指定位置周围的上下文"""
    # 查找参考文献部分的开始位置
    references_start_idx = None
    for i, paragraph in enumerate(doc.paragraphs):
        if '参考文献' in paragraph.text or 'References' in paragraph.text:
            references_start_idx = i
            break

    # 如果引用在参考文献部分，则不提取上下文
    if references_start_idx is not None and para_idx >= references_start_idx:
        return "引用出现在参考文献部分"

    paragraph = doc.paragraphs[para_idx]
    para_text = paragraph.tex

    # 获取引用在段落中的文本
    citation_text = para_text[start_pos:end_pos]

    # 提取引用周围的上下文
    context_start = max(0, start_pos - context_length)
    context_end = min(len(para_text), end_pos + context_length)
    context = para_text[context_start:context_end]

    # 如果需要更多上下文，可以考虑前后段落
    # 但要确保不包含参考文献部分
    if len(context) < context_length * 2:
        # 添加前一个段落的内容（如果存在且不为空）
        if para_idx > 0 and doc.paragraphs[para_idx-1].text.strip() and
           (references_start_idx is None or para_idx-1 < references_start_idx):
            prev_text = doc.paragraphs[para_idx-1].tex
            context = prev_text[-context_length:] + " " + contex

        # 添加后一个段落的内容（如果存在且不为空）
        if para_idx < len(doc.paragraphs) - 1 and doc.paragraphs[para_idx+1].text.strip() and
           (references_start_idx is None or para_idx+1 < references_start_idx):
            next_text = doc.paragraphs[para_idx+1].tex
            context = context + " " + next_text[:context_length]

    return contex


# 保持原有的函数以确保向后兼容性
def extract_references(docx_path: str, config: Optional[Dict] = None) -> List[str]:
    """提取文档中的作者年份格式引用（向后兼容）"""
    refs_with_context = extract_references_with_context(docx_path, config)
    return [ref['citation'] for ref in refs_with_context]


def extract_western_references(docx_path: str, config: Optional[Dict] = None) -> List[str]:
    """专门提取西式格式的引用（向后兼容）"""
    refs_with_context = extract_western_references_with_context(docx_path, config)
    return [ref['citation'] for ref in refs_with_context]


def extract_citations_from_text(text: str, config: Optional[Dict] = None) -> List[str]:
    """从给定文本中提取引用"""
    patterns = [
        # 中文格式: 作者（年份）
        r'([\u4e00-\u9fa5\w\s\.&＆,，]+?)\s*[（(](\d{4})[）)]',
        # 西文格式: Author (year)
        r'([A-Z][a-z]+(?:\s+[A-Z]\.)?(?:\s*&\s*[A-Z][a-z]+(?:\s+[A-Z]\.)?)?)\s*[（(](\d{4})[）)]',
        # et al. 格式
        r'([A-Z][a-z]+\s+et\s+al\.)[\s\u00a0]*[（(](\d{4})[）)]',
        # 等 格式
        r'([\u4e00-\u9fa5\w\s]+?)\s+等[\s\u00a0]*[（(](\d{4})[）)]',
        # 多作者格式: Johnson & Brown (2021)
        r'([A-Z][a-z]+(?:\s*&\s*[A-Z][a-z]+)+)\s*[（(](\d{4})[）)]',
        # 简单的年份格式: (Smith, 2020)
        r'[（(]([A-Z][a-z]+(?:\s*&\s*[A-Z][a-z]+)?),?\s*(\d{4})[）)]',
    ]

    references = []
    for pattern in patterns:
        matches = list(re.finditer(pattern, text))
        for match in matches:
            if len(match.groups()) >= 2:
                # 根据正则模式调整组索引
                if pattern == r'[（(]([A-Z][a-z]+(?:\s*&\s*[A-Z][a-z]+)?),?\s*(\d{4})[）)]':
                    # 这个模式中，作者在group(1)，年份在group(2)
                    author = match.group(1).strip()
                    year = match.group(2).strip()
                else:
                    # 其他模式中，作者在group(1)，年份在group(2)
                    author = match.group(1).strip()
                    year = match.group(2).strip()

                if len(author) >= 2 and year.isdigit() and len(year) == 4:
                    reference = f"{author}（{year}）"
                    references.append(reference)

    # 去重
    unique_references = list(dict.fromkeys(references))  # 保持顺序的去重

    # 使用AI优化（如果可用）
    try:
        from .ai_optimizer import optimize_citations_with_ai
        if unique_references:
            optimized = optimize_citations_with_ai(unique_references, config)
            return optimized
        else:
            return unique_references
    except ImportError:
        # 如果AI优化模块不可用，返回原始提取结果
        return unique_references


def extract_citations_from_markdown(markdown_content: str, config: Optional[Dict] = None) -> List[str]:
    """从Markdown内容中提取引用"""
    # 调用从文本提取引用的函数
    return extract_citations_from_text(markdown_content, config)