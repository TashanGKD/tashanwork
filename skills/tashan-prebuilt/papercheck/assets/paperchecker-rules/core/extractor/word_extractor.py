from collections import Counter
import re

import docx

from .base_extractor import BaseExtractor
from models.document import Document, Citation, Reference
from utils.doc_converter import DocConversionError, prepare_word_source


class WordExtractor(BaseExtractor):
    """Word文档提取器"""

    def validate_file(self, file_path: str) -> bool:
        return file_path.lower().endswith(('.docx', '.doc', '.docm'))

    def extract(self, file_path: str) -> Document:
        """提取Word文档内容"""
        prepared = prepare_word_source(file_path)
        try:
            doc = docx.Document(str(prepared.path))

            paragraphs = [paragraph.text for paragraph in doc.paragraphs]

            tables_content = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        tables_content.append(cell.text)

            references_start = self._find_reference_start(paragraphs)

            # 引文提取使用已解析的 docx 路径，避免 .doc 原始路径触发额外兼容问题
            citations = self._extract_citations(paragraphs, tables_content, str(prepared.path), references_start)
            numeric_citation_count = sum(1 for citation in citations if citation.format_type == "number")
            has_explicit_reference_indices = bool(self._collect_reference_indices(paragraphs, references_start))
            # 数字制且文末未显式编号时，按文末顺序映射，不能因为去重改变隐式编号。
            preserve_reference_slots = numeric_citation_count > 0 and not has_explicit_reference_indices
            references = self._extract_references(
                paragraphs,
                references_start,
                preserve_duplicates=preserve_reference_slots,
                preserve_reference_slots=preserve_reference_slots,
            )

            return Document(
                content=paragraphs,
                tables=tables_content,
                citations=citations,
                references=references,
                metadata={
                    "file_type": prepared.source_type,
                    "file_path": file_path,
                    "resolved_file_path": str(prepared.path),
                    "doc_conversion_tool": prepared.conversion_tool,
                },
            )
        except DocConversionError:
            raise
        finally:
            prepared.cleanup()

    def _extract_citations(self, paragraphs: list, tables_content: list, file_path: str, references_start: int = -1) -> list:
        """提取引文"""
        all_citations = []
        author_year_citations = []
        processed_author_year_texts = set()
        numeric_citation_count = 0

        try:
            from .ai_extractor import extract_references, extract_western_references
            ai_enhanced_extraction_available = True
        except ImportError:
            ai_enhanced_extraction_available = False

        body_paragraphs = paragraphs[:references_start] if references_start != -1 else paragraphs
        body_text = "\n".join(body_paragraphs)
        valid_reference_ids = self._collect_reference_indices(paragraphs, references_start)
        valid_reference_ids = valid_reference_ids or None

        for paragraph in body_paragraphs:
            if self._is_toc_entry(paragraph):
                continue

            for citation_id in self._extract_numeric_citation_ids(paragraph, valid_reference_ids):
                all_citations.append(Citation(text=f"[{citation_id}]", format_type='number', context=paragraph))
                numeric_citation_count += 1

            for citation in self._extract_author_year_citations(paragraph):
                normalized_text = self._normalize_extracted_author_year_text(citation.text)
                if not normalized_text or normalized_text in processed_author_year_texts:
                    continue
                citation.text = normalized_text
                author_year_citations.append(citation)
                processed_author_year_texts.add(normalized_text)

        for cell_text in tables_content:
            for citation_id in self._extract_numeric_citation_ids(cell_text, valid_reference_ids):
                all_citations.append(Citation(text=f"[{citation_id}]", format_type='number', context=cell_text))
                numeric_citation_count += 1

        is_numeric_dominant = numeric_citation_count >= max(5, len(author_year_citations) * 2)
        if not is_numeric_dominant:
            if ai_enhanced_extraction_available:
                try:
                    ai_config = {
                        "api_key": "your-api-key",
                        "model_name": "qwen-plus"
                    }
                    ai_author_year = extract_references(file_path, ai_config)
                    ai_western = extract_western_references(file_path, ai_config)
                    for citation_text in ai_author_year + ai_western:
                        normalized_text = self._normalize_extracted_author_year_text(citation_text)
                        if not normalized_text or normalized_text in processed_author_year_texts:
                            continue
                        if not self._is_ai_citation_supported_by_body(normalized_text, body_text):
                            continue
                        author_year_citations.append(
                            Citation(text=normalized_text, format_type='author_year', context="AI提取")
                        )
                        processed_author_year_texts.add(normalized_text)
                except Exception as e:
                    print(f"AI增强引用提取失败: {e}")

            all_citations.extend(author_year_citations)

        return all_citations

    def _extract_author_year_citations(self, paragraph: str) -> list:
        """提取作者年份格式引用"""
        citations = []
        normalized_paragraph = re.sub(
            r'([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])和([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])',
            r'\1 and \2',
            paragraph or "",
        )
        normalized_paragraph = re.sub(
            r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})\s*和\s*([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})',
            r'\1 and \2',
            normalized_paragraph,
        )
        normalized_paragraph = re.sub(
            r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})\s*等(?=\s*[（(])',
            r'\1 et al.',
            normalized_paragraph,
        )
        normalized_paragraph = re.sub(
            r'([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{1,})and([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]+)',
            r'\1 and \2',
            normalized_paragraph,
        )
        chinese_pattern = re.compile(
            r'(?P<author>(?:[\u4e00-\u9fff]{2,4}(?:\s*[、和及]\s*[\u4e00-\u9fff]{2,4})*(?:\s*等)?))\s*[（(](?P<year>(?:18|19|20)\d{2})(?P<suffix>[a-z]?)[）)]'
        )
        for match in chinese_pattern.finditer(normalized_paragraph):
            author = self._normalize_author_year_author(match.group("author"), is_chinese=True)
            if not author:
                continue
            if self._is_noise_like_cn_author_token(author):
                continue
            year = match.group("year")
            suffix = (match.group("suffix") or "").lower()
            citations.append(Citation(
                text=f"{author}（{year}{suffix}）",
                format_type='author_year',
                author=author,
                year=f"{year}{suffix}",
                context=paragraph
            ))

        english_pattern = re.compile(
            r'(?P<author>[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ\'’\-]+(?:\s+(?:&|and|et al\.?|[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ\'’\-]+)){0,6})\s*[（(](?P<year>(?:18|19|20)\d{2})(?P<suffix>[a-z]?)[）)]',
            re.IGNORECASE,
        )
        for match in english_pattern.finditer(normalized_paragraph):
            author = self._normalize_author_year_author(match.group("author"), is_chinese=False)
            if not author:
                continue
            if self._looks_like_non_author_entity(author):
                continue
            year = match.group("year")
            suffix = (match.group("suffix") or "").lower()
            citations.append(Citation(
                text=f"{author} ({year}{suffix})",
                format_type='author_year',
                author=author,
                year=f"{year}{suffix}",
                context=paragraph
            ))

        return citations

    def _extract_references(
        self,
        paragraphs: list,
        references_start: int = -1,
        preserve_duplicates: bool = False,
        preserve_reference_slots: bool = False,
    ) -> list:
        """提取参考文献（块级策略优先，兼容传统逐行提取）"""
        if references_start == -1:
            references_start = self._find_reference_start(paragraphs)

        numbered_block_refs = self._extract_references_numbered_blocks(paragraphs, references_start)
        if len(numbered_block_refs) >= 30:
            return self._remove_duplicates_and_validate(
                numbered_block_refs,
                preserve_duplicates=preserve_duplicates,
                preserve_reference_slots=preserve_reference_slots,
            )

        unnumbered_block_refs = self._extract_references_unnumbered_blocks(paragraphs, references_start)
        traditional_refs = self._extract_references_traditional(paragraphs, references_start)

        if references_start != -1 and len(unnumbered_block_refs) >= 20:
            return self._remove_duplicates_and_validate(
                numbered_block_refs + unnumbered_block_refs,
                preserve_duplicates=preserve_duplicates,
                preserve_reference_slots=preserve_reference_slots,
            )

        if references_start != -1 and max(len(unnumbered_block_refs), len(traditional_refs)) >= 10:
            return self._remove_duplicates_and_validate(
                numbered_block_refs + unnumbered_block_refs + traditional_refs,
                preserve_duplicates=preserve_duplicates,
                preserve_reference_slots=preserve_reference_slots,
            )

        return self._remove_duplicates_and_validate(
            numbered_block_refs + unnumbered_block_refs + traditional_refs,
            preserve_duplicates=preserve_duplicates,
            preserve_reference_slots=preserve_reference_slots,
        )

    def _extract_references_numbered_blocks(self, paragraphs: list, references_start: int = -1) -> list:
        """按 [n] 编号块提取参考文献"""
        references = []
        if not paragraphs:
            return references

        if references_start == -1:
            references_start = self._find_reference_start(paragraphs)
        if references_start == -1:
            references_start = int(len(paragraphs) * 0.7)

        number_start_pattern = re.compile(r'^\s*\[(\d+)\]\s*(.*)$')
        current_lines = []

        for i in range(references_start + 1, len(paragraphs)):
            text = (paragraphs[i] or "").strip()
            if not text:
                continue
            if self._is_end_marker(text):
                break

            start_match = number_start_pattern.match(text)
            if start_match:
                self._append_reference_from_lines(references, current_lines)
                number = start_match.group(1)
                remainder = start_match.group(2).strip()
                if remainder:
                    current_lines = [f"[{number}] {remainder}"]
                else:
                    current_lines = [f"[{number}]"]
                continue

            if current_lines:
                if text.startswith("#") and len(text) < 40:
                    continue
                current_lines.append(text)

        self._append_reference_from_lines(references, current_lines)
        return references

    def _extract_references_unnumbered_blocks(self, paragraphs: list, references_start: int = -1) -> list:
        """按“作者起始行”提取无编号参考文献块"""
        references = []
        if not paragraphs:
            return references

        if references_start == -1:
            references_start = self._find_reference_start(paragraphs)
        if references_start == -1:
            references_start = int(len(paragraphs) * 0.7)

        line_counts = Counter((p or "").strip() for p in paragraphs if (p or "").strip())
        current_lines = []

        for i in range(references_start + 1, len(paragraphs)):
            text = (paragraphs[i] or "").strip()
            if not text:
                continue
            if self._is_end_marker(text):
                break
            if self._is_reference_header_footer_line(paragraphs, i, text):
                continue
            if self._is_reference_noise_line(text, line_counts):
                continue

            if current_lines and self._looks_like_unnumbered_reference_start(text):
                if self._should_merge_with_current_reference(current_lines, text):
                    current_lines.append(text)
                    continue
                self._append_reference_from_lines(references, current_lines)
                current_lines = [text]
                continue

            current_lines.append(text)

        self._append_reference_from_lines(references, current_lines)
        return references

    def _extract_references_traditional(self, paragraphs: list, references_start: int = -1) -> list:
        """传统逐行提取（兼容兜底）"""
        references = []
        if references_start == -1:
            references_start = self._find_reference_start(paragraphs)

        if references_start == -1:
            return references

        for i in range(references_start + 1, len(paragraphs)):
            text = (paragraphs[i] or "").strip()
            if not text:
                continue
            if self._is_end_marker(text):
                break

            if self._is_reference_entry(text):
                from core.checker.citation_checking.reference_mapper import extract_author_year_from_reference
                extracted_author, extracted_year = extract_author_year_from_reference(text)
                references.append(Reference(text=text, author=extracted_author, year=extracted_year))

        return references

    def _append_reference_from_lines(self, references: list, lines: list):
        """将多行参考文献拼接为单条并写入结果"""
        if not lines:
            return

        merged = " ".join(part.strip() for part in lines if part and part.strip())
        merged = re.sub(r'\s+', ' ', merged).strip()
        merged = re.sub(r'\s+([,.;:])', r'\1', merged)
        merged = re.sub(r'([(\[（［])\s+', r'\1', merged)
        merged = re.split(r'(?:致\s*谢|Acknowledgements?)', merged, maxsplit=1, flags=re.IGNORECASE)[0].strip()
        if not merged:
            return

        from core.checker.citation_checking.reference_mapper import extract_author_year_from_reference

        for candidate in self._split_compound_reference_entry(merged):
            extracted_author, extracted_year = extract_author_year_from_reference(candidate)
            references.append(Reference(text=candidate, author=extracted_author, year=extracted_year))

    def _split_compound_reference_entry(self, merged: str) -> list:
        """拆分单行粘连的多条参考文献"""
        if not merged:
            return []

        split_pattern = re.compile(
            r"(?<=[。\.])\s*(?=(?:"
            r"(?:[\u4e00-\u9fff]{2,4}|[\u4e00-\u9fff]{2,20}(?:局|部|委|院|所|司|中心|集团|公司|大学|学院|学会|协会|研究院|政府|日报|新闻网|报|网|出版社))\s*[,，]\s*[\u4e00-\u9fffA-Za-z]"
            r"|[\u4e00-\u9fff·]{2,30}\s*[,，]\s*[\u4e00-\u9fffA-Za-z]"
            r"|[\u4e00-\u9fff]{2,4}\."
            r"|[\u4e00-\u9fff]{2,20}(?:局|部|委|院|所|司|中心|集团|公司|大学|学院|学会|协会|研究院|政府|日报|新闻网|报|网|出版社)\."
            r"|[A-Z]{2,}(?:\s+[A-Z]{2,}){0,5}\.(?:\s+|$)"
            r"|[A-Z][A-Za-z]+[A-Z]{2,}[A-Za-z]*\.(?:\s+|$)"
            r"|[A-Z][A-Za-z'’&\-]+(?:\s+[A-Z][A-Za-z'’&\-]+){1,8}(?:\s+\([A-Z]{2,}\))?\.(?:\s+|$)"
            r"|[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]+(?:\s+[A-Z]{1,3}(?:\.)?){0,4}\s*,\s*[A-ZÀ-ÖØ-Þ]"
            r"|[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]+(?:\s+[A-Z]{1,3}(?:\.)?){1,5}\.(?:\s+|$)"
            r"|[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]+\s+[A-Z]{1,3}\.(?=[A-Z])"
            r"))"
        )

        parts = [part.strip() for part in split_pattern.split(merged) if part and part.strip()]
        if not parts:
            return [merged]
        return self._merge_split_reference_fragments(parts)

    def _remove_duplicates_and_validate(
        self,
        references: list,
        preserve_duplicates: bool = False,
        preserve_reference_slots: bool = False,
    ) -> list:
        """去重并校验参考文献条目"""
        seen_texts = set()
        seen_text_list = []
        unique_refs = []

        for ref in references:
            normalized_text = ' '.join(ref.text.split())
            if not preserve_duplicates and normalized_text in seen_texts:
                continue
            if self._looks_like_reference_tail_fragment(normalized_text):
                if any(existing.endswith(normalized_text) and len(existing) > len(normalized_text) + 8 for existing in seen_text_list):
                    continue
            if not preserve_duplicates:
                seen_texts.add(normalized_text)
            seen_text_list.append(normalized_text)
            if preserve_reference_slots:
                if self._is_reference_entry(ref.text):
                    unique_refs.append(ref)
                continue
            if self._is_valid_reference(ref.text):
                unique_refs.append(ref)

        return unique_refs

    def _is_valid_reference(self, text: str) -> bool:
        """校验候选是否是有效参考文献条目"""
        if len(text) < 10:
            return False

        is_numbered = bool(re.match(r'^\s*\[\d+\]', text))
        has_doc_type_marker = bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', text))
        has_year = bool(re.search(r'\b(19|20)\d{2}\b', text))
        if not has_year and not is_numbered and not has_doc_type_marker:
            return False

        # 过滤明显续行碎片：如 "2015,(20):11-14."
        if (not is_numbered and not has_doc_type_marker and
                re.match(r'^\s*(19|20)\d{2}\s*[,，]\s*\(\d+\)\s*[:：]\s*\d+', text)):
            return False

        has_locator = bool(re.search(
            r'(?:\b\d+\(\d+\)\s*[:：]\s*[A-Za-z]?\d+|[:：]\s*\d+\s*[-–]\s*\d+|'
            r'\bVol\.?\s*\d+|\bNo\.?\s*\d+|\bpp?\.?\s*\d+|doi\s*[:：]|https?://)',
            text,
            re.IGNORECASE
        ))

        exclude_keywords = ['附录', '致谢', '作者简历', 'Acknowledgements', 'Appendix']
        for keyword in exclude_keywords:
            if keyword in text[:50]:
                return False

        sentence_punct_count = text.count('。') + text.count('！') + text.count('？') + text.count('；') + text.count(';')
        if len(text) >= 120 and sentence_punct_count >= 1 and not is_numbered and not has_doc_type_marker and not has_locator:
            return False

        has_plain_reference_shape = bool(re.search(r'\.\s*[^.]{2,120}\.\s*[^,，]{2,120}[,，]\s*(?:19|20)\d{2}', text))
        if not (is_numbered or has_doc_type_marker or has_locator or has_plain_reference_shape):
            return False

        return True

    def _is_reference_entry(self, text: str) -> bool:
        """判断是否为参考文献条目"""
        has_year = bool(re.search(r'\d{4}', text))
        has_doc_type_marker = bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', text))
        has_basic_structure = len(text) > 10 and ('.' in text or '[' in text)

        exclude_keywords = ['附录', '致谢', '作者简历']
        has_exclude_keyword = any(keyword in text for keyword in exclude_keywords)

        return (has_year or has_doc_type_marker) and has_basic_structure and not has_exclude_keyword

    def _is_end_marker(self, text: str) -> bool:
        """判断是否为参考文献部分结束标记"""
        end_keywords = ['附录', '致谢', '作者简历', 'Appendix', 'Acknowledgements']
        if any(keyword in text for keyword in end_keywords) and len(text) < 80:
            return True
        if re.search(r'致\s*谢', text):
            return True
        if text.strip() in {'致', '谢'}:
            return True
        return False

    def _is_toc_entry(self, text: str) -> bool:
        """判断是否为目录行（标题 + 页码）"""
        stripped = text.strip()
        return bool(re.search(r'\t\s*\d+\s*$', stripped) or re.search(r'\.{2,}\s*\d+\s*$', stripped))

    def _find_reference_start(self, paragraphs: list) -> int:
        """定位参考文献起点，优先真实段落而非目录/页眉重复标题"""
        if not paragraphs:
            return -1

        exact_candidates = []
        contains_candidates = []
        for i, paragraph in enumerate(paragraphs):
            text = (paragraph or "").strip()
            low = text.lower()
            if low in {'参考文献', '# 参考文献', 'references', '# references'}:
                exact_candidates.append(i)
            elif ('参考文献' in text or 'references' in low) and not self._is_toc_entry(text):
                contains_candidates.append(i)

        def pick_best(candidates: list) -> int:
            if not candidates:
                return -1
            scored = [(self._score_reference_start_candidate(paragraphs, idx), idx) for idx in candidates]
            best_score = max(score for score, _ in scored)
            best_indexes = [idx for score, idx in scored if score == best_score]
            return min(best_indexes)

        if exact_candidates:
            return pick_best(exact_candidates)
        if contains_candidates:
            return pick_best(contains_candidates)
        return -1

    def _score_reference_start_candidate(self, paragraphs: list, start_idx: int) -> int:
        """给参考文献起点候选打分"""
        score = 0
        window_end = min(len(paragraphs), start_idx + 200)

        for i in range(start_idx + 1, window_end):
            text = (paragraphs[i] or "").strip()
            if not text:
                continue
            if self._is_end_marker(text):
                break
            if self._is_toc_entry(text):
                continue
            if re.fullmatch(r'\d{1,4}', text):
                continue

            if re.match(r'^\s*\[\d+\]', text):
                score += 5
                continue

            if self._looks_like_unnumbered_reference_start(text):
                score += 4
            if re.search(r'\b(19|20)\d{2}\b', text):
                score += 1
            if re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', text):
                score += 2

        return score

    def _is_reference_header_footer_line(self, paragraphs: list, index: int, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False

        has_year = bool(re.search(r'\b(19|20)\d{2}\b', stripped))
        has_doc_type = bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', stripped))
        if has_year or has_doc_type:
            return False
        if any(sep in stripped for sep in [',', '，', ':', '：', '.', '。']):
            return False
        if self._looks_like_unnumbered_reference_start(stripped):
            return False

        prev_text = ""
        for j in range(index - 1, -1, -1):
            prev_text = (paragraphs[j] or "").strip()
            if prev_text:
                break

        next_text = ""
        for j in range(index + 1, len(paragraphs)):
            next_text = (paragraphs[j] or "").strip()
            if next_text:
                break

        if re.fullmatch(r'\d{1,4}', prev_text) or re.fullmatch(r'\d{1,4}', next_text):
            return len(stripped) >= 6

        return False

    def _looks_like_unnumbered_reference_start(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        if re.match(r'^\s*\[\d+\]', stripped):
            return True
        if re.match(r'^(?:[\u4e00-\u9fff]{2,4}|[\u4e00-\u9fff]{2,20}(?:局|部|委|院|所|司|中心|集团|公司|大学|学院|学会|协会|研究院|政府|日报|新闻网|报|网|出版社))\s*[,，]\s*[\u4e00-\u9fffA-Za-z]', stripped):
            return True
        if re.match(r'^[\u4e00-\u9fff·]{2,30}\s*[,，]\s*[\u4e00-\u9fffA-Za-z]', stripped):
            return True
        if re.match(r'^[\u4e00-\u9fff]{2,4}\.', stripped):
            return True
        if re.match(r'^[\u4e00-\u9fff]{2,20}(?:局|部|委|院|所|司|中心|集团|公司|大学|学院|学会|协会|研究院|政府|日报|新闻网|报|网|出版社)\.', stripped):
            return True
        if re.match(r'^[A-Z]{2,}(?:\s+[A-Z]{2,}){0,5}\.(?:\s+|$)', stripped):
            return True
        if re.match(r'^[A-Z][A-Za-z]+[A-Z]{2,}[A-Za-z]*\.(?:\s+|$)', stripped):
            return True
        if re.match(r"^[A-Z][A-Za-z'’&\-]+(?:\s+[A-Z][A-Za-z'’&\-]+){1,8}(?:\s+\([A-Z]{2,}\))?\.(?:\s+|$)", stripped):
            return True
        if re.match(r"^[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]+(?:\s+[A-Z]{1,3}(?:\.)?){0,4}\s*,\s*[A-ZÀ-ÖØ-Þ]", stripped):
            return True
        if re.match(r"^[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]+(?:\s+[A-Z]{1,3}\s*(?:\.)?){1,5}\s*\.(?:\s+|$)", stripped):
            return True
        if re.match(r'^[A-Z](?:\s+[A-Z]){1,5}\.\s+[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]', stripped):
            return True
        if re.match(r"^[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]{2,}$", stripped):
            return True
        if re.match(r'^[A-Z][A-Z\s\.-]{2,30},\s*[A-Z]', stripped):
            return True
        return False

    def _is_reference_noise_line(self, text: str, line_counts: Counter) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return True
        low = stripped.lower()
        if low in {'参考文献', '# 参考文献', 'references', '# references'}:
            return True
        if re.fullmatch(r'\d{1,4}', stripped):
            return True
        if stripped in {'致', '谢'}:
            return True
        if self._is_toc_entry(stripped):
            return True

        has_year = bool(re.search(r'\b(19|20)\d{2}\b', stripped))
        has_doc_type = bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', stripped))
        if line_counts.get(stripped, 0) >= 2 and not has_year and not has_doc_type and len(stripped) >= 8:
            return True

        return False

    def _extract_numeric_citation_ids(self, text: str, valid_reference_ids=None) -> list:
        """从文本中提取数字引用并展开区间/并列，如 [11-14], [32, 33]"""
        if self._is_noise_like_numeric_context(text):
            return []

        ids = []
        token_pattern = re.compile(
            r'[\[［〔【](\d+(?:\s*[-–]\s*\d+)?(?:\s*[,，;；、]\s*\d+(?:\s*[-–]\s*\d+)?)*)[\]］〕】]'
        )
        for match in token_pattern.finditer(text):
            raw = match.group(1)
            parts = [part.strip() for part in re.split(r'[,，;；、]', raw)]
            for part in parts:
                range_match = re.match(r'^(\d+)\s*[-–]\s*(\d+)$', part)
                if range_match:
                    start = int(range_match.group(1))
                    end = int(range_match.group(2))
                    if start <= end:
                        for idx in range(start, end + 1):
                            if idx <= 0:
                                continue
                            if valid_reference_ids is not None and idx not in valid_reference_ids:
                                continue
                            ids.append(idx)
                    continue

                if part.isdigit():
                    idx = int(part)
                    if idx <= 0:
                        continue
                    if valid_reference_ids is not None and idx not in valid_reference_ids:
                        continue
                    ids.append(idx)

        return ids

    def _collect_reference_indices(self, paragraphs: list, references_start: int = -1) -> set:
        if not paragraphs:
            return set()
        if references_start == -1:
            references_start = self._find_reference_start(paragraphs)
        if references_start == -1:
            return set()

        number_start_pattern = re.compile(r'^\s*[\[［〔【](\d+)[\]］〕】]')
        indices = set()
        for i in range(references_start + 1, len(paragraphs)):
            text = (paragraphs[i] or "").strip()
            if not text:
                continue
            if self._is_end_marker(text):
                break
            match = number_start_pattern.match(text)
            if not match:
                continue
            number = int(match.group(1))
            if number > 0:
                indices.add(number)
        return indices

    def _is_noise_like_numeric_context(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return True

        if re.fullmatch(
            r'[\[［〔【]\d+(?:\s*[-–]\s*\d+)?(?:\s*[,，;；、]\s*\d+(?:\s*[-–]\s*\d+)?)*[\]］〕】]',
            stripped,
        ) and len(stripped) <= 16:
            return True

        math_operator_count = len(re.findall(r'[=+\-*/<>≤≥∑Σ∫]', stripped))
        alpha_count = len(re.findall(r'[A-Za-z\u4e00-\u9fff]', stripped))
        if math_operator_count >= 2 and alpha_count <= 3:
            return True

        return False

    def _normalize_author_year_author(self, author_text: str, is_chinese: bool) -> str:
        text = re.sub(r'<br\s*/?>', ' ', str(author_text or ''), flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text).strip(" ,，;；:：。.")
        if not text:
            return ""

        leading_connectors = (
            r"(?:据|由|与|及|并|和|在|对|关于|针对|根据|基于|依据|参考|采用|使用|比如|例如|如|以及|其中|对于|随着|通过|本文|文中|本研究|尽管|回应了|印证了|指出|认为|提出|发现|显示|强调|揭示|呼吁|证实)"
        )
        text = re.sub(rf"^(?:{leading_connectors})+", "", text).strip(" ,，;；:：。.")
        text = re.sub(r'^(?:本研究借鉴|借鉴|尝试回应|该理论由|它起源于|该理论的核心观点由)', '', text).strip(" ,，;；:：。.")
        text = re.sub(r'([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])和([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])', r'\1 and \2', text)
        text = re.sub(r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})\s*和\s*([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})', r'\1 and \2', text)
        text = re.sub(
            r'([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{1,})and([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]+)',
            r'\1 and \2',
            text,
        )
        text = re.sub(r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})\s*等$', r'\1 et al.', text)
        text = re.sub(r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})etal\.?', r'\1 et al.', text)
        text = re.sub(r'\s*&\s*', ' & ', text)

        if is_chinese:
            match = re.search(r'([\u4e00-\u9fff]{2,4}(?:\s*[、和及]\s*[\u4e00-\u9fff]{2,4})*(?:\s*等)?)$', text)
            if match:
                text = match.group(1)
        else:
            match = re.search(
                r'([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ\'’\-]+(?:\s+(?:&|and|et al\.?|[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ\'’\-]+)){0,6})$',
                text,
                flags=re.IGNORECASE,
            )
            if match:
                text = match.group(1)
            text = re.sub(r'\betal\b', 'et al.', text, flags=re.IGNORECASE)

        return re.sub(r'\s+', ' ', text).strip()

    def _normalize_extracted_author_year_text(self, citation_text: str) -> str:
        text = (citation_text or "").strip()
        if text.startswith("[AUTH:") and text.endswith("]"):
            text = text[6:-1].strip()
        text = re.sub(r'<br\\s*/?>', ' ', text, flags=re.IGNORECASE)
        text = re.sub(r'\s+', ' ', text)
        text = text.replace("（", "(").replace("）", ")")
        text = re.sub(r'([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])和([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ])', r'\1 and \2', text)
        text = re.sub(
            r'([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{1,})and([A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]+)',
            r'\1 and \2',
            text,
        )
        text = re.sub(r'(?i)([A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]{2,})etal\.?', r'\1 et al.', text)
        text = re.sub(r'\s*&\s*', ' & ', text)

        pair_pattern = re.compile(
            r'(?P<author>[^()]{1,120}?)\s*[\\(](?P<year>(?:18|19|20)\d{2})(?P<suffix>[a-z]?)[\\)]',
            re.IGNORECASE,
        )
        matches = list(pair_pattern.finditer(text))
        if not matches:
            return ""
        match = matches[-1]
        author = (match.group("author") or "").strip()
        year = (match.group("year") or "").strip()
        suffix = (match.group("suffix") or "").lower()
        if not author or not year:
            return ""

        has_english = bool(re.search(r'[A-Za-z]', author))
        normalized_author = self._normalize_author_year_author(author, is_chinese=not has_english)
        if not normalized_author:
            return ""
        if (not has_english and self._is_noise_like_cn_author_token(normalized_author)) or (
            has_english and self._looks_like_non_author_entity(normalized_author)
        ):
            return ""
        if has_english:
            return f"{normalized_author} ({year}{suffix})"
        return f"{normalized_author}（{year}{suffix}）"

    def _is_ai_citation_supported_by_body(self, citation_text: str, body_text: str) -> bool:
        from core.checker.citation_checking.reference_mapper import extract_author_year_from_citation

        author, year = extract_author_year_from_citation(citation_text or "")
        if not author or not year:
            return False

        normalized_author = re.sub(r"\s+", " ", author).strip()
        normalized_year = str(year).strip().lower()
        if not normalized_author or not normalized_year:
            return False

        compact_body = re.sub(r"\s+", "", body_text or "")
        normalized_body = re.sub(r"\s+", " ", body_text or "").replace("（", "(").replace("）", ")")
        year_pattern = re.escape(normalized_year)

        if re.search(r"[\u4e00-\u9fff]", normalized_author):
            compact_author = re.sub(r"\s+", "", normalized_author)
            variants = {
                compact_author,
                compact_author.replace("等", ""),
            }
            for variant in variants:
                if not variant:
                    continue
                explicit_pattern = rf"(?<![\u4e00-\u9fffA-Za-z0-9]){re.escape(variant)}(?:等)?[（(]{year_pattern}[）)]"
                parenthetical_pattern = rf"[（(]{re.escape(variant)}(?:等)?[，,]?\s*{year_pattern}[）)]"
                for match in re.finditer(explicit_pattern, compact_body):
                    if self._is_likely_trailing_author_fragment(compact_author, compact_body, match.start()):
                        continue
                    return True
                for match in re.finditer(parenthetical_pattern, compact_body):
                    if self._is_likely_trailing_author_fragment(compact_author, compact_body, match.start() + 1):
                        continue
                    return True
            return False

        for variant in self._build_english_author_lookup_variants(normalized_author):
            variant_pattern = re.escape(variant).replace(r"\ ", r"\s+")
            explicit_pattern = rf"(?<![A-Za-zÀ-ÖØ-öø-ÿĀ-ſ]){variant_pattern}\s*[（(]\s*{year_pattern}\s*[）)]"
            parenthetical_pattern = rf"[（(]\s*{variant_pattern}\s*[,，]\s*{year_pattern}\s*[）)]"
            for match in re.finditer(explicit_pattern, normalized_body, flags=re.IGNORECASE):
                if self._is_likely_trailing_author_fragment(normalized_author, normalized_body, match.start()):
                    continue
                return True
            for match in re.finditer(parenthetical_pattern, normalized_body, flags=re.IGNORECASE):
                if self._is_likely_trailing_author_fragment(normalized_author, normalized_body, match.start() + 1):
                    continue
                return True
        return False

    def _is_likely_trailing_author_fragment(self, author: str, body_text: str, match_start: int) -> bool:
        normalized_author = re.sub(r"\s+", " ", author or "").strip()
        if not normalized_author or not self._is_single_author_marker(normalized_author):
            return False
        if match_start <= 0:
            return False
        prefix = (body_text or "")[:match_start]
        prefix = re.sub(r"\s+$", "", prefix)
        if not prefix:
            return False
        if re.search(r"(?:[、，,&/]|and|with|和|及)\s*$", prefix, flags=re.IGNORECASE):
            return True
        return False

    def _is_single_author_marker(self, author: str) -> bool:
        normalized = re.sub(r"\s+", " ", str(author or "")).strip()
        if not normalized:
            return False
        if re.search(r"[\u4e00-\u9fff]", normalized):
            compact = re.sub(r"\s+", "", normalized)
            if any(mark in compact for mark in ("、", "和", "及", "等")):
                return False
            return len(compact) >= 2

        low = normalized.lower()
        if " et al" in low or " and " in low or "&" in normalized:
            return False
        tokens = [token for token in normalized.split(" ") if token]
        return len(tokens) <= 2

    def _build_english_author_lookup_variants(self, author: str) -> set:
        normalized = re.sub(r"\s+", " ", author or "").strip()
        if not normalized:
            return set()

        variants = {
            normalized,
            normalized.replace(" and ", " & "),
            normalized.replace(" & ", " and "),
            re.sub(r"(?i)\bet\s+al\.?\b", "et al.", normalized),
            re.sub(r"(?i)\bet\s+al\.?\b", "et al", normalized),
        }
        return {variant.strip() for variant in variants if variant and variant.strip()}

    def _is_noise_like_cn_author_token(self, author: str) -> bool:
        compact = re.sub(r"\s+", "", str(author or ""))
        if not compact:
            return True
        if compact in {"学者", "等", "等学者", "研究者", "受访者", "作者"}:
            return True
        if any(token in compact for token in ("年报", "报告")):
            return True
        if compact.startswith("等") and len(compact) <= 4:
            return True
        return False

    def _looks_like_non_author_entity(self, author: str) -> bool:
        normalized = re.sub(r"\s+", " ", str(author or "")).strip()
        if not normalized:
            return True
        lowered = normalized.lower()
        product_like_tokens = {
            "ipad",
            "iphone",
            "airpods",
            "apple watch",
            "applewatch",
            "ipod",
            "itunes",
            "macbook",
        }
        if lowered in product_like_tokens:
            return True
        if re.search(r"\d", lowered):
            return True
        return False

    def _looks_like_standalone_latin_surname(self, text: str) -> bool:
        stripped = (text or "").strip()
        return bool(re.match(r"^[A-ZÀ-ÖØ-Þ][A-Za-zÀ-ÖØ-öø-ÿĀ-ſ'’\-]{2,}$", stripped))

    def _looks_like_latin_surname_continuation(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        if re.match(r'^[A-Z](?:\.|\b)(?:\s+[A-Z](?:\.|\b)){0,6}\s+', stripped):
            return True
        if re.match(r'^[A-Z][a-z]+(?:\s+[a-z][a-z]+){1,6}', stripped):
            return True
        return False

    def _should_merge_with_current_reference(self, current_lines: list, next_line: str) -> bool:
        if not current_lines:
            return False
        last_line = (current_lines[-1] or "").strip()
        if not self._looks_like_standalone_latin_surname(last_line):
            return False
        return self._looks_like_latin_surname_continuation(next_line)

    def _merge_split_reference_fragments(self, parts: list) -> list:
        merged = []
        i = 0
        while i < len(parts):
            current = parts[i]
            if i + 1 < len(parts):
                nxt = parts[i + 1]
                if self._looks_like_english_author_lead_fragment(current) and self._looks_like_doc_type_continuation(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_chinese_author_lead_fragment(current) and self._looks_like_doc_type_continuation(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_standalone_chinese_name(current) and self._looks_like_doc_type_continuation(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_standalone_chinese_name(current) and not self._looks_like_unnumbered_reference_start(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._should_merge_conference_continuation(current, nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_doc_type_tail(current) and self._looks_like_reference_source_suffix(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_doc_type_tail(current) and not self._looks_like_unnumbered_reference_start(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
                if self._looks_like_standalone_latin_surname(current) and self._looks_like_latin_surname_continuation(nxt):
                    merged.append(f"{current} {nxt}".strip())
                    i += 2
                    continue
            merged.append(current)
            i += 1
        return merged

    def _looks_like_doc_type_tail(self, text: str) -> bool:
        stripped = (text or "").strip()
        return bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]\.\s*$', stripped))

    def _looks_like_doc_type_continuation(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        has_doc_type = bool(re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', stripped))
        has_year = bool(re.search(r'\b(19|20)\d{2}\b', stripped))
        return has_doc_type and has_year

    def _looks_like_english_author_lead_fragment(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped or not stripped.endswith('.'):
            return False
        if not re.search(r'[A-Za-z]', stripped):
            return False
        if re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', stripped):
            return False
        if re.search(r'\b(19|20)\d{2}\b', stripped):
            return False
        if ',' not in stripped:
            return False
        return True

    def _looks_like_chinese_author_lead_fragment(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        return bool(re.match(r'^[\u4e00-\u9fff]{2,4}(?:\s*[,，]\s*[\u4e00-\u9fff]{2,4}){1,8}\.$', stripped))

    def _looks_like_reference_source_suffix(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        if re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', stripped):
            return False
        if not re.search(r'\b(19|20)\d{2}\b', stripped):
            return False
        if re.match(r'^[\u4e00-\u9fff]{2,12}[\.：:]\s*[\u4e00-\u9fffA-Za-z]', stripped):
            return True
        if re.match(r"^[A-Z][A-Za-z'’&\-\s]{2,60}:\s*[A-ZÀ-ÖØ-Þ]", stripped):
            return True
        if re.match(r"^[A-Z][A-Za-z'’&\-\s]{2,80},\s*(?:19|20)\d{2}", stripped):
            return True
        return False

    def _looks_like_standalone_chinese_name(self, text: str) -> bool:
        stripped = (text or "").strip()
        return bool(re.match(r'^[\u4e00-\u9fff]{2,4}\.$', stripped))

    def _should_merge_conference_continuation(self, current: str, nxt: str) -> bool:
        current_text = (current or "").strip()
        next_text = (nxt or "").strip()
        if not current_text or not next_text:
            return False

        if not re.search(r'[\[［]\s*C(?:/[A-Za-z]+)?\s*[\]］]\.\s*$', current_text, flags=re.IGNORECASE):
            return False

        if not re.match(
            r'^[\u4e00-\u9fff]{2,20}(?:学会|协会|大学|学院|研究院|研究所|出版社|集团|公司|中心|委员会|委|部|局|司)\.',
            next_text,
        ):
            return False

        # 若后段很快出现新的文献类型标识，更可能是独立条目，不合并。
        if re.search(r'[\[［]\s*[A-Za-z]+(?:/[A-Za-z]+)?\s*[\]］]', next_text[:40]):
            return False
        return True

    def _looks_like_reference_tail_fragment(self, text: str) -> bool:
        stripped = (text or "").strip()
        if not stripped:
            return False
        if re.match(r'^[A-Z](?:\.)\s+[A-Z][a-z]', stripped):
            return True
        if re.match(r'^(?:19|20)\d{2}\s*[,，]\s*\(\d+\)\s*[:：]\s*\d+', stripped):
            return True
        return False

    def _extract_doi(self, text: str) -> str:
        """提取DOI"""
        doi_pattern = r'doi:\s*([^\s,.;\)]+)|DOI:\s*([^\s,.;\)]+)'
        match = re.search(doi_pattern, text, re.IGNORECASE)
        if match:
            return match.group(1) or match.group(2)
        return None

    def _extract_url(self, text: str) -> str:
        """提取URL"""
        url_pattern = r'https?://[^\s,.;\)]+'
        match = re.search(url_pattern, text)
        return match.group(0) if match else None
