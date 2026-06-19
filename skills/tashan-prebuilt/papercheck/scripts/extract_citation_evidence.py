#!/usr/bin/env python
"""Extract citation evidence from a DOCX without using any model API."""

from __future__ import annotations

import argparse
import json
import re
from pathlib import Path

import docx


CITATION_RE = re.compile(r"\[(\d+)(?:-(\d+))?\]")


def paragraph_texts(document: docx.Document) -> list[str]:
    texts = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    texts.append(text)
    return texts


def reference_start(texts: list[str]) -> int:
    for i, text in enumerate(texts):
        compact = re.sub(r"\s+", "", text).lower()
        if "参考文献" in compact or compact == "references":
            return i
    return len(texts)


def expand_citation(match: re.Match[str]) -> list[int]:
    start = int(match.group(1))
    end = int(match.group(2) or start)
    if end < start or end - start > 200:
        return [start]
    return list(range(start, end + 1))


def extract_reference_number(text: str) -> int | None:
    match = re.match(r"^\s*\[(\d+)\]\s*(.+)", text)
    if match:
        return int(match.group(1))
    return None


def context_for(citation_number: int, body: list[str], window: int) -> dict:
    citation = f"[{citation_number}]"
    hits = []
    for i, text in enumerate(body):
        direct = citation in text
        ranged = False
        for match in CITATION_RE.finditer(text):
            if citation_number in expand_citation(match):
                ranged = True
                break
        if direct or ranged:
            start = max(0, i - window)
            end = min(len(body), i + window + 1)
            hits.append(
                {
                    "paragraph_index": i,
                    "context": " ".join(body[start:end]),
                    "matched_text": text,
                    "match_type": "direct" if direct else "range",
                }
            )
    return {"citation": citation, "contexts": hits}


def build_evidence(docx_path: Path, window: int, max_contexts: int) -> dict:
    document = docx.Document(str(docx_path))
    texts = paragraph_texts(document)
    ref_start = reference_start(texts)
    body = texts[:ref_start]
    ref_lines = texts[ref_start + 1 :] if ref_start < len(texts) else []

    cited_numbers: set[int] = set()
    citation_mentions = []
    for i, text in enumerate(body):
        for match in CITATION_RE.finditer(text):
            numbers = expand_citation(match)
            cited_numbers.update(numbers)
            citation_mentions.append(
                {
                    "paragraph_index": i,
                    "raw": match.group(0),
                    "expanded": numbers,
                    "text": text,
                }
            )

    references = {}
    unnumbered_references = []
    for text in ref_lines:
        number = extract_reference_number(text)
        if number is None:
            if text and "参考文献" not in text:
                unnumbered_references.append(text)
            continue
        references[number] = text

    items = []
    for number in sorted(cited_numbers):
        ctx = context_for(number, body, window)
        contexts = ctx["contexts"][:max_contexts]
        items.append(
            {
                "citation": f"[{number}]",
                "number": number,
                "reference": references.get(number),
                "reference_found": number in references,
                "contexts": contexts,
                "context_count": len(ctx["contexts"]),
                "needs_model_review": bool(contexts and number in references),
            }
        )

    reference_numbers = set(references)
    return {
        "source_docx": str(docx_path),
        "paragraph_count": len(texts),
        "body_paragraph_count": len(body),
        "reference_count": len(references),
        "citation_count": len(cited_numbers),
        "citation_mentions": citation_mentions,
        "missing_citations": [f"[{n}]" for n in sorted(cited_numbers - reference_numbers)],
        "unused_references": [f"[{n}]" for n in sorted(reference_numbers - cited_numbers)],
        "unnumbered_reference_lines": unnumbered_references,
        "items": items,
    }


def main() -> int:
    parser = argparse.ArgumentParser(description="Extract citation evidence from DOCX.")
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path)
    parser.add_argument("--window", type=int, default=1)
    parser.add_argument("--max-contexts", type=int, default=3)
    args = parser.parse_args()

    if not args.docx.exists():
        raise SystemExit(f"Input file not found: {args.docx}")
    if args.docx.suffix.lower() != ".docx":
        raise SystemExit("Only .docx input is supported by this evidence extractor.")

    evidence = build_evidence(args.docx, args.window, args.max_contexts)
    out = args.out or args.docx.with_suffix(".citation-evidence.json")
    out.write_text(json.dumps(evidence, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps({"ok": True, "out": str(out), "citation_count": evidence["citation_count"], "reference_count": evidence["reference_count"]}, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
